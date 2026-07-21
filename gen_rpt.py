"""RAG Chatbot Technical Architecture Report Generator"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime, os, sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15
for lv in range(1, 4):
    h = doc.styles[f"Heading {lv}"]
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    h.font.name = "Calibri"

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x1E,0x1E,0x1E)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    shd = p._element.get_or_add_pPr()
    bg = shd.makeelement(qn("w:shd"), {qn("w:val"):"clear", qn("w:color"):"auto", qn("w:fill"):"F5F5F5"})
    shd.append(bg)

def tbl(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs: r.bold = True
    for rd in rows:
        cells = t.add_row().cells
        for i, v in enumerate(rd): cells[i].text = str(v)
    doc.add_paragraph()

def diag(title, lines):
    doc.add_paragraph(title, style="Heading 3")
    p = doc.add_paragraph()
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"; r.font.size = Pt(9)

def save():
    OUT = "RAG_Chatbot_Technical_Architecture_Report.docx"
    doc.save(OUT)
    print(f"Report saved: {OUT}")
    print(f"Full path: {os.path.abspath(OUT)}")

# TITLE PAGE
for _ in range(4): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("RAG Chatbot\nTechnical Architecture Report")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("DotSquares AI - Multi-Agent RAG Assistant")
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Date: {datetime.date.today().strftime('%d %B %Y')}\nVersion: 0.1.0\nStatus: Current Codebase Snapshot").font.size = Pt(11)
doc.add_page_break()

# TABLE OF CONTENTS
doc.add_heading("Table of Contents", level=1)
toc_items = ["1. Project Overview","2. Complete Folder Structure","3. Backend Architecture","4. Chat Flow","5. Retrieval Architecture","6. Current Retrieval Pipeline","7. Database Schema","8. API Endpoints","9. Frontend Architecture","10. Authentication Flow","11. Current Features","12. AI Components","13. Project Dependencies","14. Request Lifecycle","15. Current Limitations","16. Existing Design Patterns","17. Final Project Summary"]
for item in toc_items:
    p = doc.add_paragraph(item); p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# SECTION 1
doc.add_heading("1. Project Overview", level=1)
doc.add_heading("1.1 Purpose", level=2)
doc.add_paragraph("The RAG Chatbot is a full-stack Retrieval-Augmented Generation application that answers questions about DotSquares company policies. It retrieves relevant passages from a PDF document and generates natural-language answers using a locally-hosted LLM.")
doc.add_heading("1.2 Problem Solved", level=2)
doc.add_paragraph("Employees need quick, accurate answers about company policies (leave, WFH, code of conduct, HR rules, security) without manually searching through a long PDF. The chatbot provides conversational access to this knowledge base.")
doc.add_heading("1.3 Overall Architecture", level=2)
doc.add_paragraph("Client-server architecture. Frontend (Next.js) communicates with backend (FastAPI) via REST APIs through Next.js proxy rewrites. Backend orchestrates authentication, chat session management, and the RAG pipeline. RAG pipeline uses hybrid retrieval (vector search via ChromaDB + keyword search via BM25), followed by LLM-based reranking and answer generation via locally-running Ollama model.")
diag("System Architecture", [
    "+------------------+    REST/JSON     +----------------------------------+",
    "|                  | <---------------> |        FastAPI Backend            |",
    "|   Next.js        | (via proxy rewrite)|   +----------------------------+ |",
    "|   Frontend       |                   |   |  Auth Layer                | |",
    "|  (port 3000)     |                   |   |  X-API-Key + JWT Bearer    | |",
    "+------------------+                   |   +----------------------------+ |",
    "                                       |   |  Chat Orchestration        | |",
    "                                       |   |  Session/Message CRUD      | |",
    "                                       |   |  ChatbotService.ask()      | |",
    "                                       |   +----------------------------+ |",
    "                                       |   |  RAG Pipeline              | |",
    "                                       |   |  QueryRewriter             | |",
    "                                       |   |  HybridRetriever           | |",
    "                                       |   |  Reranker (LLM)            | |",
    "                                       |   +----------------------------+ |",
    "                                       +--------------+-------------------+",
    "                                                      |",
    "                        +------------------------------+------------------+",
    "                        |                              |                  |",
    "                   +----v-----+  +--------------+  +---v-----------+",
    "                   |PostgreSQL|  |  ChromaDB     |  |   Ollama      |",
    "                   | (Users,  |  |  (Vector      |  |  qwen3:4b     |",
    "                   | Sessions,|  |   Store)      |  |  nomic-embed  |",
    "                   | Messages)|  +--------------+  +---------------+",
    "                   +----------+",
])
doc.add_heading("1.4 Tech Stack Summary", level=2)
tbl(["Layer","Technology","Details"],[
    ["Backend","FastAPI v0.138.2","Async Python web framework with Uvicorn v0.49.0"],
    ["Database","PostgreSQL","via psycopg2-binary 2.9.12, SQLAlchemy 2.0.51 ORM"],
    ["Migrations","Alembic v1.18.5","4 migration files for schema evolution"],
    ["AI Framework","LangChain v1.3.11","langchain-core 1.4.8, langchain-ollama 1.1.0"],
    ["Agent Framework","LangGraph v1.2.7","Defined but not wired into main flow"],
    ["LLM Runtime","Ollama v0.6.2","Local inference, no cloud API calls"],
    ["LLM Model","qwen3:4b","Temperature=0, timeout=300s"],
    ["Embedding Model","nomic-embed-text","Local, via OllamaEmbeddings"],
    ["Vector DB","ChromaDB v1.5.9","Persisted at backend/chatbot/chroma_db/"],
    ["Keyword Search","BM25Okapi","rank-bm25 v0.2.2 with regex tokenizer"],
    ["Frontend","Next.js v14.2.18","App Router with React 18.3.1, TypeScript ~5.6.2"],
    ["Markdown","react-markdown v10.1.0","Renders AI responses with formatting"],
])

# SECTION 2
doc.add_heading("2. Complete Folder Structure", level=1)
code("""rag_chatbot/
├── README.md
├── .gitignore
├── backend/
│   ├── main.py                          # FastAPI entry point, CORS, routers, global exception handler
│   ├── requirements.txt                 # 162 pinned Python dependencies
│   ├── alembic.ini                      # Alembic migration config
│   ├── .env / .env.example              # Environment variables
│   ├── config/
│   │   └── environment.py               # Loads env vars: DB URL, JWT secret, API key, media path
│   ├── db/
│   │   ├── __init__.py                  # Re-exports Base, engine, SessionLocal, get_db
│   │   └── db_session.py               # SQLAlchemy engine, sessionmaker, get_db dependency
│   ├── shared/
│   │   ├── __init__.py                  # Re-exports get_logger, log_error, custom responses
│   │   ├── logging/__init__.py          # get_logger(), log_error() utilities
│   │   ├── auth/
│   │   │   ├── __init__.py              # Re-exports JWT + API key functions
│   │   │   ├── create_access_token.py   # create_jwt_token(), verify_jwt_token()
│   │   │   ├── jwt_auth_function.py     # FastAPI Bearer dependency jwt_auth()
│   │   │   └── x_api_auth.py           # FastAPI X-API-Key dependency x_api_auth()
│   │   └── utils/
│   │       ├── __init__.py
│   │       └── custom_responses.py      # custom_error_response(), custom_response()
│   ├── src/routes/
│   │   ├── users/
│   │   │   ├── router.py                # Users APIRouter mounting sub-routers
│   │   │   ├── models/
│   │   │   │   ├── user_ref.py          # UserRef ORM (users table)
│   │   │   │   └── user_profiles.py     # UserProfile ORM (user_profile table)
│   │   │   ├── schemas/
│   │   │   │   ├── login_request.py     # LoginRequest Pydantic model
│   │   │   │   └── profile_out.py       # ProfileOut Pydantic model
│   │   │   ├── controllers/
│   │   │   │   ├── get_user_by_email.py
│   │   │   │   ├── hash_password.py     # Argon2 hashing
│   │   │   │   ├── verify_password.py   # Argon2 verification
│   │   │   │   └── validate_password.py # Password strength rules
│   │   │   └── routers/
│   │   │       ├── signup.py            # POST /users/api/signup
│   │   │       ├── login.py             # POST /users/api/login
│   │   │       ├── get_user_profile.py  # GET /users/profile-data
│   │   │       └── update_user_profile.py # PATCH /users/update-profile
│   │   └── chatbot/
│   │       ├── router.py                # Chatbot APIRouter mounting sessions+messages
│   │       ├── models/
│   │       │   ├── chat_session.py      # ChatSession ORM
│   │       │   └── chat_message.py      # ChatMessage ORM (JSONB sources)
│   │       ├── schemas/
│   │       │   ├── chat_request.py      # ChatRequest(content: str)
│   │       │   ├── chat_response.py     # ChatMessageOut
│   │       │   ├── session_request.py   # SessionRequest, SessionUpdateRequest
│   │       │   ├── session_response.py  # ChatSessionOut
│   │       │   └── chat_schema.py       # Standalone schemas (for unmounted chat.py)
│   │       ├── crud/
│   │       │   ├── chat_session.py      # Session CRUD + title derivation
│   │       │   └── chat_message.py      # Message CRUD
│   │       ├── controllers/
│   │       │   ├── session_controller.py # send_chat_message() MAIN ORCHESTRATION
│   │       │   └── chat_controller.py   # Standalone (for unmounted chat.py)
│   │       └── routers/
│   │           ├── sessions.py          # Session CRUD endpoints
│   │           ├── messages.py          # Message endpoints
│   │           └── chat.py              # Standalone /chatbot/chat (NOT mounted)
│   ├── chatbot/                         # RAG ENGINE (standalone package)
│   │   ├── config.py                    # LLM_MODEL, EMBEDDING_MODEL, TOP_K constants
│   │   ├── prompts.py                   # SYSTEM_PROMPT with {history}, {context}, {question}
│   │   ├── chatbot_service.py           # ChatbotService singleton
│   │   ├── vector_store.py             # ChromaDB singleton
│   │   ├── bm25_store.py              # BM25Okapi retriever singleton
│   │   ├── hybrid_retriever.py         # Merges vector + BM25 with SHA-256 dedup
│   │   ├── tokenizer.py               # Regex tokenizer for BM25
│   │   ├── document_loader.py          # PDF load + text split
│   │   ├── ingest.py                   # Standalone ingestion script
│   │   ├── router.py                   # Standalone endpoint (NOT mounted)
│   │   ├── retrieval/
│   │   │   ├── pipeline.py             # RetrievalPipeline orchestrator
│   │   │   ├── query_rewriter.py       # Heuristic + LLM query rewriting
│   │   │   └── reranker.py            # LLM-based document reranking
│   │   ├── agents/
│   │   │   ├── supervisor.py           # Keyword-based routing (rag/web)
│   │   │   ├── rag_agent.py            # Wraps ChatbotService.ask()
│   │   │   └── web_agent.py            # STUB - returns "not implemented"
│   │   ├── graph/
│   │   │   ├── state.py                # GraphState TypedDict
│   │   │   └── graph.py               # LangGraph StateGraph (NOT wired in)
│   │   ├── nodes/
│   │   │   ├── supervisor_node.py
│   │   │   ├── rag_node.py
│   │   │   └── web_node.py
│   │   └── memory/
│   │       ├── history_loader.py       # Loads last 10 messages from DB
│   │       ├── history_selector.py     # Takes last N messages (unused)
│   │       ├── prompt_context.py       # Builds "Role: Content" text
│   │       ├── memory_manager.py       # BROKEN - calls nonexistent methods
│   │       └── conversation_summary.py # STUB - returns None
│   ├── alembic/                         # DB migration scripts (4 versions)
│   ├── test_graph.py                    # Manual LangGraph test
│   └── chatbot/data/
│       └── Dotsquares_Company_Policy.pdf  # Source PDF
└── frontend/
    ├── package.json / tsconfig.json
    ├── next.config.mjs                  # Proxy rewrites to backend
    ├── .env.local / .env.example
    ├── scripts/
    │   ├── check-node.cjs              # Node version check
    │   └── dev-fresh.cjs               # Fresh dev startup
    ├── public/images/                   # Logo/icon assets
    ├── app/
    │   ├── globals.css                  # All application CSS
    │   ├── layout.tsx                   # Root layout (fonts, metadata)
    │   ├── page.tsx                     # Root redirect logic
    │   ├── login/page.tsx              # Login form
    │   ├── signup/page.tsx             # Signup form
    │   └── (app)/
    │       ├── layout.tsx               # AppShell wrapper
    │       ├── dashboard/page.tsx       # Main chat interface
    │       ├── history/page.tsx         # Chat history list
    │       └── profile/page.tsx         # Profile management
    ├── components/
    │   ├── AppShell.tsx                 # Auth guard + sidebar
    │   ├── ApiKeySetupNotice.tsx        # API key guidance
    │   ├── DotSquaresBrandLogo.tsx      # Brand logo
    │   ├── PasswordInput.tsx            # Password toggle field
    │   └── chat/
    │       ├── ChatBubble.tsx           # Message rendering
    │       ├── ChatComposer.tsx         # Input + send button
    │       ├── AgentStepTimeline.tsx    # Pipeline visualization
    │       └── SourcePanel.tsx          # Source attribution panel
    └── lib/
        ├── api.ts                       # HTTP client, error handling
        ├── chat-api.ts                  # Chat API functions
        ├── auth-storage.ts             # localStorage auth management
        └── types/chat.ts               # TypeScript types + formatters""")

doc.add_heading("2.1 Folder Interaction Map", level=2)
doc.add_paragraph("Backend has three layers: (1) src/routes/ for HTTP handling (routers, controllers, schemas, CRUD, models), (2) chatbot/ for AI engine (retrieval, LLM, agents, memory), (3) shared/+db/+config/ for cross-cutting concerns (auth, logging, database, config). Frontend has two layers: (1) app/ for pages and routing via Next.js App Router, (2) components/+lib/ for reusable UI components and API utilities. The chatbot/ package is self-contained and can operate independently of the HTTP layer.")

# SECTION 3
doc.add_heading("3. Backend Architecture", level=1)

doc.add_heading("3.1 Entry Point - main.py", level=2)
doc.add_paragraph("File: backend/main.py")
doc.add_paragraph("Framework: FastAPI v0.138.2. CORS allows all origins, methods, headers. Static files served at /media from MEDIA_PATH. Global exception handler catches all unhandled exceptions and returns JSON {success, status, message, error}. Two routers mounted under /users prefix: users_router (login, signup, profile) and chatbot_router (sessions, messages).")
code("app.include_router(users_router,  prefix='/users')\napp.include_router(chatbot_router, prefix='/users')")
doc.add_paragraph("Resulting URL structure: /users/api/login, /users/api/signup, /users/profile-data, /users/update-profile (from users_router); /users/api/chat/sessions/*, /users/api/chat/sessions/{id}/messages (from chatbot_router).")

doc.add_heading("3.2 Authentication Module", level=2)
doc.add_paragraph("Files: backend/shared/auth/create_access_token.py, jwt_auth_function.py, x_api_auth.py; backend/src/routes/users/controllers/hash_password.py, verify_password.py, validate_password.py")
doc.add_paragraph("Two-layer authentication:")
doc.add_paragraph("Layer 1 - X-API-Key (x_api_auth.py): Every signup/login request must include X-API-Key header matching X_API_KEY env var. Returns 403 on mismatch.")
doc.add_paragraph("Layer 2 - JWT Bearer (jwt_auth_function.py): Protected endpoints use HTTPBearer scheme. Token verified via PyJWT (HS256). Returns 401 on invalid/expired. Payload: {user_id, email, roles}.")
doc.add_paragraph("Token lifecycle: create_jwt_token(data, expiry_minutes, remember_me) encodes JWT. Default expiry 24h; remember_me=True gives 30 days. verify_jwt_token(token) decodes and validates.")
doc.add_paragraph("Password handling: Argon2 hashing via passlib CryptContext(schemes=['argon2']). validate_password() checks: length>=8, uppercase, lowercase, digit, special char (@$!%*?).")

doc.add_heading("3.3 Users Module", level=2)
doc.add_paragraph("File: backend/src/routes/users/router.py - mounts signup, login, get_user_profile, update_user_profile sub-routers.")
doc.add_paragraph("ORM Models:")
doc.add_paragraph("  UserRef (users table) - backend/src/routes/users/models/user_ref.py: id, email, password_hash, display_name, dob, created_by_admin, status, email_verified, otp, last_login, otp_created_at, is_active, is_deleted, created_at, updated_at, roles. Relationship: user_profile (one-to-one).")
doc.add_paragraph("  UserProfile (user_profile table) - backend/src/routes/users/models/user_profiles.py: id, user_id (FK->users), profile_picture_url, portfolio (JSONB), watchlist (JSONB), is_active, is_deleted, created_at, updated_at.")
doc.add_paragraph("Controllers: get_user_by_email(db, email) queries UserRef; hash_password/verify_password use argon2; validate_password checks strength rules.")
doc.add_paragraph("Routers: POST /users/api/signup (Form: display_name, email, password; X-API-Key auth; creates UserRef + UserProfile), POST /users/api/login (JSON: email, password, remember_me; X-API-Key auth; returns JWT + user data), GET /users/profile-data (JWT auth; returns ProfileOut), PATCH /users/update-profile (JWT auth; Form: display_name?, dob?, profile_picture?; saves image to media/profile_pictures/)")

doc.add_heading("3.4 Chat Module", level=2)
doc.add_paragraph("File: backend/src/routes/chatbot/router.py - mounts sessions_router and messages_router under /api/chat.")
doc.add_paragraph("ORM Models:")
doc.add_paragraph("  ChatSession (chat_sessions) - backend/src/routes/chatbot/models/chat_session.py: id, user_id (FK->users), title, is_active, is_deleted, domain_key, created_at, updated_at. Relationships: user (UserRef), messages (ChatMessage[] with cascade delete-orphan).")
doc.add_paragraph("  ChatMessage (chat_messages) - backend/src/routes/chatbot/models/chat_message.py: id, session_id (FK->chat_sessions), role, content, sources (JSONB), created_at. Relationship: session (ChatSession).")
doc.add_paragraph("CRUD (backend/src/routes/chatbot/crud/): chat_session.py has create_session, get_sessions (with message count via func.count), get_session, update_session, delete_session (soft delete sets is_deleted=True), backfill_new_chat_titles, _derive_title (truncates to 50 chars). chat_message.py has create_message, get_messages (ordered by created_at ASC).")
doc.add_paragraph("Controllers: session_controller.py send_chat_message(db, session_id, user_id, content) is the MAIN ORCHESTRATION FUNCTION. It: (1) verifies session belongs to user, (2) creates user message in DB, (3) auto-titles 'New Chat' sessions, (4) loads history via HistoryLoader(db).load(session_id), (5) calls ChatbotService.ask(question=content, history=history), (6) creates assistant message with answer + sources. chat_controller.py is standalone (for unmounted chat.py endpoint).")

doc.add_heading("3.5 RAG Engine (chatbot/ package)", level=2)
doc.add_paragraph("File: backend/chatbot/config.py - Constants: BASE_DIR, DATA_PATH (PDF path), CHROMA_DB_PATH, LLM_MODEL='qwen3:4b', EMBEDDING_MODEL='nomic-embed-text', OLLAMA_TIMEOUT_SECONDS=300, VECTOR_TOP_K=10, BM25_TOP_K=10, FINAL_TOP_K=5.")
doc.add_paragraph("File: backend/chatbot/prompts.py - SYSTEM_PROMPT template with {history}, {context}, {question} variables. Rules: answer only from context, use history for follow-ups, format with Markdown, never mention internals (PDFs, vector DBs, etc.). Tone: professional, helpful, concise.")
doc.add_paragraph("File: backend/chatbot/chatbot_service.py - ChatbotService is a Singleton (class-level _instance + _initialized). Creates RetrievalPipeline, ChatOllama(model=qwen3:4b, temperature=0, timeout=300s), ChatPromptTemplate(SYSTEM_PROMPT), chain = prompt | llm. Method ask(question, history=None) returns {success, answer, sources: {chunks, agent_steps, history_summary, rewritten_query}}.")
doc.add_paragraph("File: backend/chatbot/vector_store.py - Module-level singleton. get_retriever(k=10) returns ChromaDB retriever with OllamaEmbeddings(nomic-embed-text). Persists at chatbot/chroma_db/.")
doc.add_paragraph("File: backend/chatbot/bm25_store.py - Module-level singleton. BM25Retriever loads chunks from document_loader.get_chunks(), tokenizes, creates BM25Okapi. invoke(query, k) returns ranked Document list.")
doc.add_paragraph("File: backend/chatbot/hybrid_retriever.py - HybridRetriever creates vector_retriever + bm25_retriever. invoke(query) merges both, deduplicates by SHA-256 hash of page_content.")
doc.add_paragraph("File: backend/chatbot/tokenizer.py - tokenize(text) returns lowercase alphanumeric tokens via regex.")
doc.add_paragraph("File: backend/chatbot/document_loader.py - load_documents() uses PyPDFLoader; split_documents() uses RecursiveCharacterTextSplitter(1000, 200); get_chunks() combines both.")
doc.add_paragraph("File: backend/chatbot/ingest.py - Standalone script: load_pdf -> split_documents -> Chroma.from_documents. Run: python -m chatbot.ingest")

doc.add_heading("3.6 Retrieval Pipeline", level=2)
doc.add_paragraph("File: backend/chatbot/retrieval/pipeline.py - RetrievalPipeline creates QueryRewriter, HybridRetriever, Reranker. retrieve(query): (1) rewriter.rewrite(query), (2) hybrid.invoke(rewritten_query), (3) reranker.rerank(rewritten_query, docs).")
doc.add_paragraph("File: backend/chatbot/retrieval/query_rewriter.py - QueryRewriter uses ChatOllama(qwen3:4b). needs_rewrite(query) returns True if <=2 words, known abbreviation (wfh/pto/hr/pf/esi/lop/ot/salary/leave/holiday/notice period), or no question words. rewrite(query) sends to LLM if needed.")
doc.add_paragraph("File: backend/chatbot/retrieval/reranker.py - Reranker uses ChatOllama(qwen3:4b). rerank(query, docs) sends all docs to LLM, asks for JSON array of indexes by relevance, returns top FINAL_TOP_K (5). Falls back to original order on parse failure.")

doc.add_heading("3.7 Multi-Agent Architecture (LangGraph)", level=2)
doc.add_paragraph("File: backend/chatbot/graph/state.py - GraphState(TypedDict): question, history, route, answer, sources, success, error.")
doc.add_paragraph("File: backend/chatbot/graph/graph.py - LangGraph StateGraph with nodes: supervisor, rag, web. Entry: supervisor. Conditional edges: supervisor routes to 'rag' or 'web' based on state['route']. Both connect to END. Compiled: graph = builder.compile(). STATUS: Defined but NOT wired into main request flow.")
doc.add_paragraph("File: backend/chatbot/agents/supervisor.py - SupervisorAgent.route(question) checks for keywords (today, latest, news, current, recent, google, internet, search, 2026). Returns 'web' if found, else 'rag'.")
doc.add_paragraph("File: backend/chatbot/agents/rag_agent.py - RagAgent.run(question, history) calls ChatbotService.ask().")
doc.add_paragraph("File: backend/chatbot/agents/web_agent.py - WebAgent.search(question) returns 'Web Search is not implemented yet.' (STUB).")
doc.add_paragraph("File: backend/chatbot/nodes/*.py - Module-level singleton agents. Node functions take state dict, call agent, update state.")
doc.add_paragraph("File: backend/test_graph.py - Manual test: graph.invoke({question: 'What is leave policy?', history: []})")

doc.add_heading("3.8 Memory Module", level=2)
doc.add_paragraph("File: backend/chatbot/memory/history_loader.py - HistoryLoader.__init__(db: Session). load(session_id, limit=10) queries ChatMessage, orders by created_at DESC, reverses to chronological, returns [{role, content}, ...]. STATUS: ACTIVE - called by session_controller.send_chat_message().")
doc.add_paragraph("File: backend/chatbot/memory/prompt_context.py - PromptContext.build(history) formats as 'Role: Content' lines, or 'No previous conversation.' if empty. STATUS: ACTIVE - called by ChatbotService.ask().")
doc.add_paragraph("File: backend/chatbot/memory/history_selector.py - HistorySelector.select(messages, max_messages=6) returns messages[-max_messages:]. STATUS: NOT USED.")
doc.add_paragraph("File: backend/chatbot/memory/memory_manager.py - MemoryManager orchestrates loader + selector + builder. BROKEN: calls nonexistent load_recent_messages() (actual method is load()); imports nonexistent PromptContextBuilder (actual class is PromptContext). STATUS: NOT USED.")
doc.add_paragraph("File: backend/chatbot/memory/conversation_summary.py - ConversationSummarizer.summarize(history) returns None (STUB). STATUS: NOT USED.")

doc.add_heading("3.9 Database Module", level=2)
doc.add_paragraph("File: backend/db/db_session.py - engine = create_engine(DATABASE_URL, pool_size=5, max_overflow=10, pool_pre_ping=True). SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine). Base = declarative_base(). get_db() is a FastAPI dependency: opens session, commits on success, rolls back on exception, closes in finally block.")

doc.add_heading("3.10 Configuration Module", level=2)
doc.add_paragraph("File: backend/config/environment.py - Loads .env via python-dotenv. Variables: BASE_URL, PG_HOST, PG_PORT, PG_USER, PG_PASSWORD, PG_DATABASE, DATABASE_URL (constructed), SECRET_KEY, JWT_ALGORITHM, X_API_KEY, TOKEN_EXPIRY_HOURS, MEDIA_PATH, APP_NAME.")

doc.add_heading("3.11 Shared Utilities", level=2)
doc.add_paragraph("File: backend/shared/logging/__init__.py - get_logger(name, level=DEBUG) creates logger with stdout handler. log_error(logger, message, **extra) logs error with exc_info=True.")
doc.add_paragraph("File: backend/shared/utils/custom_responses.py - custom_error_response(message, error, status_code) returns JSONResponse {success: False}. custom_response(message, data, status_code) returns dict {success: True}.")

# SECTION 4
doc.add_heading("4. Chat Flow", level=1)
doc.add_paragraph("Step-by-step flow when a user sends a message:")
diag("Complete Chat Flow", [
    "1. USER types message in ChatComposer",
    "   File: frontend/components/chat/ChatComposer.tsx",
    "   Calls onSubmit -> sendQuestion(e) in dashboard/page.tsx",
    "",
    "2. FRONTEND sends API request",
    "   File: frontend/lib/chat-api.ts sendChatMessage(sessionId, content)",
    "   POST /users/api/chat/sessions/{id}/messages",
    "   Headers: X-API-Key + Authorization: Bearer {token}",
    "   Body: {content: 'user message'}",
    "   Next.js proxy rewrites /users/* -> backend (next.config.mjs)",
    "",
    "3. ROUTER receives request",
    "   File: backend/src/routes/chatbot/routers/messages.py send_message()",
    "   Validates JWT via Depends(jwt_auth)",
    "   Extracts user_id from decoded token",
    "",
    "4. CONTROLLER orchestrates",
    "   File: backend/src/routes/chatbot/controllers/session_controller.py",
    "   Function: send_chat_message(db, session_id, user_id, content)",
    "     a) get_session() - verify session belongs to user",
    "     b) create_message(role='user', content=content) - persist user message",
    "     c) _derive_title() - auto-title if 'New Chat'",
    "     d) HistoryLoader(db).load(session_id) - load last 10 messages",
    "     e) ChatbotService().ask(question=content, history=history)",
    "",
    "5. CHATBOT SERVICE processes",
    "   File: backend/chatbot/chatbot_service.py ask(question, history)",
    "     a) self.pipeline.retrieve(question) - retrieval pipeline",
    "     b) PromptContext.build(history) - format history as text",
    "     c) self.chain.invoke({history, context, question}) - LLM generation",
    "",
    "6. RETRIEVAL PIPELINE executes",
    "   File: backend/chatbot/retrieval/pipeline.py retrieve(query)",
    "     a) QueryRewriter.rewrite(query) -> rewritten_query",
    "     b) HybridRetriever.invoke(rewritten_query) -> merged docs",
    "     c) Reranker.rerank(rewritten_query, docs) -> top 5 docs",
    "",
    "7. LLM GENERATES ANSWER",
    "   qwen3:4b via Ollama (ChatOllama)",
    "   Prompt: SYSTEM_PROMPT with {history}, {context}, {question}",
    "   Returns response.content (answer text)",
    "",
    "8. CONTROLLER persists response",
    "   create_message(role='assistant', content=answer, sources=sources)",
    "   Returns {user_message, assistant_message}",
    "",
    "9. FRONTEND updates UI",
    "   Replaces placeholder bubble with actual response",
    "   Displays answer with Markdown rendering (ChatBubble)",
    "   Shows SourcePanel with source chunks",
])

# SECTION 5
doc.add_heading("5. Retrieval Architecture", level=1)

components = [
    ("5.1 ingest.py - Data Ingestion", [
        ["Purpose", "Standalone script to create ChromaDB from PDF"],
        ["File", "backend/chatbot/ingest.py"],
        ["Functions", "load_pdf(), split_documents(docs), create_vector_db(chunks), main()"],
        ["Input", "backend/chatbot/data/Dotsquares_Company_Policy.pdf"],
        ["Output", "backend/chatbot/chroma_db/ directory with vector index"],
        ["Dependencies", "PyPDFLoader, RecursiveCharacterTextSplitter, OllamaEmbeddings, Chroma"],
        ["Chunk Config", "chunk_size=1000, chunk_overlap=200"],
        ["Run Command", "python -m chatbot.ingest"],
    ]),
    ("5.2 document_loader.py - PDF Loading", [
        ["Purpose", "Load PDF and split into chunks for retrieval"],
        ["File", "backend/chatbot/document_loader.py"],
        ["Functions", "load_documents(), split_documents(documents), get_chunks()"],
        ["Input", "DATA_PATH (Dotsquares_Company_Policy.pdf)"],
        ["Output", "List[Document] with page_content and metadata"],
        ["Loader", "PyPDFLoader"],
        ["Splitter", "RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)"],
        ["Used By", "BM25Retriever (loads chunks at init time)"],
    ]),
    ("5.3 vector_store.py - ChromaDB Vector Store", [
        ["Purpose", "Provide vector similarity search via ChromaDB"],
        ["File", "backend/chatbot/vector_store.py"],
        ["Functions", "get_retriever(k=10)"],
        ["Storage", "backend/chatbot/chroma_db/ (persist_directory)"],
        ["Embeddings", "OllamaEmbeddings(model='nomic-embed-text')"],
        ["Pattern", "Module-level singleton (_vector_db global)"],
        ["Output", "ChromaDB retriever with search_kwargs={'k': k}"],
        ["Used By", "HybridRetriever"],
    ]),
    ("5.4 bm25_store.py - BM25 Keyword Retrieval", [
        ["Purpose", "Provide keyword-based retrieval using BM25Okapi"],
        ["File", "backend/chatbot/bm25_store.py"],
        ["Class", "BM25Retriever"],
        ["Functions", "__init__(), invoke(query, k)"],
        ["Tokenizer", "backend/chatbot/tokenizer.py - regex lowercase + strip non-alphanum"],
        ["Data Source", "document_loader.get_chunks() (loads from PDF)"],
        ["Pattern", "Module-level singleton (_bm25 global)"],
        ["Output", "List[Document] ranked by BM25 score"],
        ["Used By", "HybridRetriever"],
    ]),
    ("5.5 hybrid_retriever.py - Hybrid Retrieval", [
        ["Purpose", "Merge vector + BM25 results with deduplication"],
        ["File", "backend/chatbot/hybrid_retriever.py"],
        ["Class", "HybridRetriever"],
        ["Functions", "__init__(), invoke(query)"],
        ["Vector K", "VECTOR_TOP_K (10)"],
        ["BM25 K", "BM25_TOP_K (10)"],
        ["Dedup", "SHA-256 hash of page_content"],
        ["Output", "List[Document] - merged, deduplicated"],
        ["Used By", "RetrievalPipeline"],
    ]),
    ("5.6 retrieval/pipeline.py - Pipeline Orchestrator", [
        ["Purpose", "Orchestrate the full retrieval pipeline"],
        ["File", "backend/chatbot/retrieval/pipeline.py"],
        ["Class", "RetrievalPipeline"],
        ["Functions", "__init__(), retrieve(query)"],
        ["Pipeline", "QueryRewriter -> HybridRetriever -> Reranker"],
        ["Input", "User query string"],
        ["Output", "List[Document] - reranked, top FINAL_TOP_K (5)"],
        ["Used By", "ChatbotService.ask()"],
    ]),
    ("5.7 retrieval/query_rewriter.py - Query Rewriting", [
        ["Purpose", "Expand/rewrite user query for better retrieval"],
        ["File", "backend/chatbot/retrieval/query_rewriter.py"],
        ["Class", "QueryRewriter"],
        ["Functions", "__init__(), needs_rewrite(query), rewrite(query)"],
        ["Heuristic", "Rewrite if <=2 words, known abbreviation, or no question words"],
        ["Abbreviations", "wfh, pto, hr, pf, esi, lop, ot, salary, leave, holiday, notice period"],
        ["LLM", "qwen3:4b via ChatOllama"],
        ["Output", "Rewritten query string (or original if no rewrite needed)"],
        ["Used By", "RetrievalPipeline"],
    ]),
    ("5.8 retrieval/reranker.py - LLM Reranking", [
        ["Purpose", "Rerank retrieved documents by relevance using LLM"],
        ["File", "backend/chatbot/retrieval/reranker.py"],
        ["Class", "Reranker"],
        ["Functions", "__init__(), rerank(query, docs)"],
        ["Method", "Sends docs to LLM, asks for JSON array of indexes by relevance"],
        ["LLM", "qwen3:4b via ChatOllama"],
        ["Output", "List[Document] - top FINAL_TOP_K (5)"],
        ["Fallback", "Returns original order if JSON parsing fails"],
        ["Used By", "RetrievalPipeline"],
    ]),
    ("5.9 chatbot_service.py - Service Orchestrator", [
        ["Purpose", "Orchestrate retrieval + LLM answer generation"],
        ["File", "backend/chatbot/chatbot_service.py"],
        ["Class", "ChatbotService (Singleton)"],
        ["Functions", "__init__(), ask(question, history=None)"],
        ["LLM", "ChatOllama(model='qwen3:4b', temperature=0, timeout=300s)"],
        ["Prompt", "SYSTEM_PROMPT with {history}, {context}, {question}"],
        ["Chain", "prompt | llm (LangChain LCEL)"],
        ["Output", "{success, answer, sources: {chunks, agent_steps, ...}}"],
        ["Used By", "session_controller.send_chat_message()"],
    ]),
]

for title, rows in components:
    doc.add_heading(title, level=2)
    tbl(["Attribute", "Detail"], rows)

# SECTION 6
doc.add_heading("6. Current Retrieval Pipeline", level=1)
diag("Retrieval Pipeline Diagram", [
    "User Query: 'What is the leave policy?'",
    "        |",
    "        v",
    "+-------------------------------------------------------+",
    "|  QueryRewriter  (retrieval/query_rewriter.py)         |",
    "|                                                       |",
    "|  needs_rewrite(query) check:                          |",
    "|    - len(query.split()) <= 2?  -> True (rewrite)     |",
    "|    - Known abbreviation?       -> True (rewrite)     |",
    "|    - Has question words?       -> False (skip)       |",
    "|                                                       |",
    "|  If rewrite: Send to qwen3:4b LLM                    |",
    "|  Prompt: 'Rewrite for company policy retrieval'       |",
    "|  Returns: expanded query string                       |",
    "+---------------------------+---------------------------+",
    "                            |",
    "                            v",
    "+-------------------------------------------------------+",
    "|  HybridRetriever  (hybrid_retriever.py)               |",
    "|                                                       |",
    "|  +-----------------+    +----------------------+       |",
    "|  | Vector Search    |    | BM25 Keyword Search   |       |",
    "|  | (ChromaDB)       |    | (BM25Okapi)           |       |",
    "|  | top 10 docs      |    | top 10 docs           |       |",
    "|  +--------+--------+    +----------+-----------+       |",
    "|           |                        |                   |",
    "|           +----------+-------------+                   |",
    "|                      v                                |",
    "|           Merge + SHA-256 Dedup                       |",
    "|           (up to ~20 docs)                            |",
    "+---------------------------+---------------------------+",
    "                            |",
    "                            v",
    "+-------------------------------------------------------+",
    "|  Reranker  (retrieval/reranker.py)                    |",
    "|                                                       |",
    "|  1. Format all docs with index numbers                |",
    "|  2. Send to qwen3:4b LLM                              |",
    "|     Prompt: 'Return JSON array of indexes by rel.'    |",
    "|  3. Parse JSON response -> reorder docs               |",
    "|  4. Return top FINAL_TOP_K (5) documents              |",
    "|  5. Fallback: original order if parse fails           |",
    "+---------------------------+---------------------------+",
    "                            |",
    "                            v",
    "Ranked Documents (top 5) -> Context for LLM Answer",
])
doc.add_paragraph("Classes: RetrievalPipeline, QueryRewriter, HybridRetriever, Reranker, BM25Retriever, get_retriever(), tokenize(), get_chunks(). Files: retrieval/pipeline.py, retrieval/query_rewriter.py, retrieval/reranker.py, hybrid_retriever.py, bm25_store.py, vector_store.py, tokenizer.py, document_loader.py.")

# SECTION 7
doc.add_heading("7. Frontend Architecture", level=1)

doc.add_heading("7.1 Project Structure", level=2)
doc.add_paragraph("Framework: Next.js 14.2.18 (App Router), React 18.3.1, TypeScript.")
doc.add_paragraph("Directory layout: app/(app)/ for authenticated pages, components/ for shared UI, lib/ for API clients and utilities.")
code("""frontend/
  app/
    page.tsx              # Root: redirect to /dashboard or /login
    layout.tsx            # Root layout: fonts (DM Sans, Outfit)
    login/page.tsx        # Login form
    signup/page.tsx       # Signup form
    (app)/
      layout.tsx          # AppShell wrapper
      dashboard/page.tsx  # Main chat interface
      history/page.tsx    # Session history list
      profile/page.tsx    # Profile editor
  components/
    AppShell.tsx           # Sidebar + content layout + auth guard
    DotSquaresBrandLogo.tsx
    PasswordInput.tsx
    ApiKeySetupNotice.tsx
    chat/
      ChatBubble.tsx      # Message bubble with markdown
      ChatComposer.tsx    # Auto-resizing textarea + send
      AgentStepTimeline.tsx
      SourcePanel.tsx
  lib/
    api.ts                # API utility functions
    chat-api.ts           # Chat/session API calls
    auth-storage.ts       # Token + session storage helpers
    types/chat.ts         # TypeScript types + constants""")

doc.add_heading("7.2 Auth & API Pattern", level=2)
doc.add_paragraph("Auth flow: signup (/signup) -> login (/login) -> stores token + user snapshot in localStorage via setAuthSession(). AppShell runs jwtAuthCheck() on mount; redirects to /login if expired.")
doc.add_paragraph("API calls: lib/chat-api.ts provides typed functions (createChatSession, listChatSessions, getChatMessages, sendChatMessage, deleteSession). lib/api.ts provides apiJson() wrapper that adds X-API-Key and Authorization headers, and apiErrorMessage() for error parsing.")
doc.add_paragraph("Proxy: next.config.mjs rewrites /users/* and /media/* to backend via the 'rewrites' config, so frontend calls hit Next.js which proxies to FastAPI.")

doc.add_heading("7.3 Dashboard Page", level=2)
doc.add_paragraph("File: app/(app)/dashboard/page.tsx - Main chat interface.")
doc.add_paragraph("Key state: sessions[], currentSessionId, messages[], isSending, error.")
doc.add_paragraph("Message flow:")
doc.add_paragraph("  1. User types in ChatComposer, onSubmit -> sendQuestion()")
doc.add_paragraph("  2. createChatSession() if no currentSessionId")
doc.add_paragraph("  3. Optimistic UI: adds user message + placeholder assistant message to state")
doc.add_paragraph("  4. sendChatMessage(sessionId, content) POST to backend")
doc.add_paragraph("  5. On success: replaces placeholder with real assistant message")
doc.add_paragraph("  6. backfillSessionTitles() to update 'New Chat' titles")
doc.add_paragraph("Key components rendered: ChatComposer, ChatBubble (with SourcePanel, AgentStepTimeline), ApiKeySetupNotice.")
doc.add_paragraph("Suggested prompts: 'What is the leave policy?', 'What is the salary structure?', 'What is the work from home policy?'")

# SECTION 8
doc.add_heading("8. Frontend - Backend Integration", level=1)

doc.add_heading("8.1 Request Flow Diagram", level=2)
diag("Frontend to Backend Request Flow", [
    "React Component (dashboard/page.tsx)",
    "    |",
    "    | calls sendChatMessage(sessionId, content)",
    "    v",
    "lib/chat-api.ts",
    "    |",
    "    | apiJson('POST', url, body)",
    "    v",
    "lib/api.ts apiJson()",
    "    |",
    "    | Adds headers:",
    "    |   X-API-Key: {publicApiKey}",
    "    |   Authorization: Bearer {accessToken}",
    "    |",
    "    v",
    "fetch() to /users/api/chat/sessions/{id}/messages",
    "    |",
    "    v",
    "next.config.mjs proxy rewrite",
    "    /users/* -> backend (http://localhost:8000)",
    "    |",
    "    v",
    "FastAPI Router: backend/src/routes/chatbot/routers/messages.py",
    "    |",
    "    | Depends(jwt_auth) validates JWT",
    "    v",
    "Controller: session_controller.send_chat_message()",
    "    |",
    "    | HistoryLoader -> ChatbotService.ask() -> ChatMessage DB write",
    "    v",
    "Response: {user_message, assistant_message}",
    "    |",
    "    v",
    "Frontend: unwrapApiData(response), update messages[] state",
])

doc.add_heading("8.2 API Endpoints Used", level=2)
tbl(["Method", "Endpoint", "Auth", "Purpose"], [
    ["POST", "/users/api/signup", "X-API-Key", "Register new user (Form data)"],
    ["POST", "/users/api/login", "X-API-Key", "Login, returns JWT (JSON)"],
    ["GET", "/users/profile-data", "JWT Bearer", "Get user profile"],
    ["PATCH", "/users/update-profile", "JWT Bearer", "Update profile (Form + image)"],
    ["POST", "/users/api/chat/sessions", "JWT Bearer", "Create new chat session"],
    ["GET", "/users/api/chat/sessions", "JWT Bearer", "List all sessions (with message count)"],
    ["PATCH", "/users/api/chat/sessions/{id}", "JWT Bearer", "Update session title"],
    ["DELETE", "/users/api/chat/sessions/{id}", "JWT Bearer", "Soft-delete session"],
    ["POST", "/users/api/chat/sessions/{id}/messages", "JWT Bearer", "Send message, get AI response"],
    ["GET", "/users/api/chat/sessions/{id}/messages", "JWT Bearer", "Load messages for a session"],
])

doc.add_heading("8.3 State Management", level=2)
doc.add_paragraph("No global state library (Redux, Zustand, etc.). All state is local via React useState/useEffect in page components.")
doc.add_paragraph("Persistent state: localStorage used for auth token (via auth-storage.ts), current session ID. Session list fetched fresh on each navigation to /history.")

doc.add_heading("8.4 Error Handling", level=2)
doc.add_paragraph("API errors: apiErrorMessage(err) in lib/api.ts extracts error message from response or returns fallback string. 403 errors trigger isForbiddenApiKeyError() check and show ApiKeySetupNotice component.")
doc.add_paragraph("Chat errors: error state in dashboard/page.tsx; error message displayed in UI. Failed sends result in placeholder messages being removed from state.")

# SECTION 9
doc.add_heading("9. Database Schema", level=1)

doc.add_heading("9.1 Tables & Relationships", level=2)
diag("Entity Relationship Diagram", [
    "+----------------+       +----------------+       +----------------+",
    "|    users       |       | user_profile   |       |  chat_sessions |",
    "+----------------+       +----------------+       +----------------+",
    "| id (PK)        |<----->| id (PK)        |       | id (PK)        |",
    "| email          |  1:1  | user_id (FK)   |       | user_id (FK)   |<-----> users.id",
    "| password_hash  |       | profile_pic... |       | title          |",
    "| display_name   |       | portfolio (J)  |       | is_active      |",
    "| dob            |       | watchlist (J)  |       | is_deleted     |",
    "| status         |       | is_active      |       | domain_key     |",
    "| email_verified |       | is_deleted     |       | created_at     |",
    "| otp            |       | created_at     |       | updated_at     |",
    "| last_login     |       | updated_at     |       +-------+--------+",
    "| otp_created_at |       +----------------+               |",
    "| is_active      |                                       | 1:N",
    "| is_deleted     |       +----------------+               |",
    "| created_at     |       | chat_messages  |               |",
    "| updated_at     |       +----------------+<--------------+",
    "| roles          |       | id (PK)        |",
    "+----------------+       | session_id(FK) |",
    "                           | role           |",
    "                           | content        |",
    "                           | sources (JSONB)|",
    "                           | created_at     |",
    "                           +----------------+",
])

doc.add_heading("9.2 Column Details", level=2)
tbl(["Table", "Column", "Type", "Constraints"], [
    ["users", "id", "integer", "PK, auto-increment"],
    ["users", "email", "varchar", "unique, not null"],
    ["users", "password_hash", "varchar", "not null (argon2)"],
    ["users", "display_name", "varchar", "nullable"],
    ["users", "dob", "date", "nullable"],
    ["users", "status", "enum", "pending, active, suspended"],
    ["users", "email_verified", "boolean", "default false"],
    ["users", "otp", "varchar", "nullable"],
    ["users", "last_login", "timestamp", "nullable"],
    ["users", "otp_created_at", "timestamp", "nullable"],
    ["users", "is_active", "boolean", "default true"],
    ["users", "is_deleted", "boolean", "default false (soft delete)"],
    ["users", "created_by_admin", "boolean", "default false"],
    ["users", "created_at", "timestamp", "server default"],
    ["users", "updated_at", "timestamp", "server default"],
    ["users", "roles", "array", "nullable"],
    ["user_profile", "id", "integer", "PK, auto-increment"],
    ["user_profile", "user_id", "integer", "FK->users.id, unique"],
    ["user_profile", "profile_picture_url", "varchar", "nullable"],
    ["user_profile", "portfolio", "jsonb", "nullable"],
    ["user_profile", "watchlist", "jsonb", "nullable"],
    ["user_profile", "is_active", "boolean", "default true"],
    ["user_profile", "is_deleted", "boolean", "default false"],
    ["user_profile", "created_at", "timestamp", "server default"],
    ["user_profile", "updated_at", "timestamp", "server default"],
    ["chat_sessions", "id", "uuid", "PK, auto-generated"],
    ["chat_sessions", "user_id", "uuid", "FK->users.id, indexed"],
    ["chat_sessions", "title", "varchar", "nullable, default 'New Chat'"],
    ["chat_sessions", "is_active", "boolean", "default true"],
    ["chat_sessions", "is_deleted", "boolean", "default false"],
    ["chat_sessions", "domain_key", "varchar", "nullable"],
    ["chat_sessions", "created_at", "timestamp", "server default"],
    ["chat_sessions", "updated_at", "timestamp", "server default"],
    ["chat_messages", "id", "uuid", "PK, auto-generated"],
    ["chat_messages", "session_id", "uuid", "FK->chat_sessions.id, indexed"],
    ["chat_messages", "role", "enum", "user, assistant"],
    ["chat_messages", "content", "text", "not null"],
    ["chat_messages", "sources", "jsonb", "nullable"],
    ["chat_messages", "created_at", "timestamp", "server default"],
])

doc.add_heading("9.3 Indexes", level=2)
doc.add_paragraph("  chat_sessions.user_id: indexed (used for listing user's sessions)")
doc.add_paragraph("  chat_messages.session_id: indexed (used for loading session messages)")
doc.add_paragraph("  users.email: unique constraint (used for login lookup)")

doc.add_heading("9.4 Migrations", level=2)
doc.add_paragraph("Alembic is configured (alembic.ini + alembic/ directory) but no explicit migration files were found in the current codebase. Tables are likely created via Base.metadata.create_all() or manual migration outside the codebase.")

# SECTION 10
doc.add_heading("10. Authentication & Authorization", level=1)

doc.add_heading("10.1 Authentication Layers", level=2)
tbl(["Layer", "Mechanism", "File", "Scope"], [
    ["Layer 1", "X-API-Key header", "shared/auth/x_api_auth.py", "Signup + Login endpoints only"],
    ["Layer 2", "JWT Bearer token", "shared/auth/jwt_auth_function.py", "All protected endpoints"],
    ["Password", "Argon2 hashing", "users/controllers/hash_password.py", "Signup + Login"],
])

doc.add_heading("10.2 JWT Token Lifecycle", level=2)
diag("JWT Token Lifecycle", [
    "LOGIN",
    "  |",
    "  v",
    "create_jwt_token(data={user_id, email, roles}, expiry_minutes, remember_me)",
    "  |",
    "  | PyJWT.encode() with HS256 + SECRET_KEY",
    "  v",
    "Token stored in localStorage (auth-storage.ts setAuthSession())",
    "  |",
    "  v",
    "PROTECTED REQUEST",
    "  |",
    "  | Authorization: Bearer {token}",
    "  v",
    "jwt_auth() -> Depends(HTTPBearer())",
    "  |",
    "  | decode_jwt_token(token) verifies HS256 signature + expiry",
    "  v",
    "request.state.decoded_token = {user_id, email, roles}",
    "  |",
    "  v",
    "CONTROLLER extracts user_id from decoded_token['user_id']",
])

doc.add_heading("10.3 Token Expiry", level=2)
tbl(["Scenario", "Expiry", "Setting"], [
    ["Default (remember_me=False)", "24 hours", "TOKEN_EXPIRY_HOURS=24 in environment.py"],
    ["Remember me (remember_me=True)", "30 days", "expiry_minutes = 60*24*30 (hardcoded)"],
])

doc.add_heading("10.4 Password Validation Rules", level=2)
doc.add_paragraph("File: backend/src/routes/users/controllers/validate_password.py")
doc.add_paragraph("Function: validate_password(password) -> bool")
tbl(["Rule", "Check"], [
    ["Minimum length", "len(password) >= 8"],
    ["Uppercase letter", "any(c.isupper() for c in password)"],
    ["Lowercase letter", "any(c.islower() for c in password)"],
    ["Digit", "any(c.isdigit() for c in password)"],
    ["Special character", "any(c in '@$!%*?' for c in password)"],
])

# SECTION 11
doc.add_heading("11. Configuration & Environment", level=1)

doc.add_heading("11.1 Backend Configuration", level=2)
tbl(["Variable", "Source", "Default", "Purpose"], [
    ["DATABASE_URL", "environment.py", "postgresql://...", "Full SQLAlchemy connection string"],
    ["PG_HOST", ".env", "127.0.0.1", "PostgreSQL host"],
    ["PG_PORT", ".env", "5432", "PostgreSQL port"],
    ["PG_USER", ".env", "postgres", "PostgreSQL user"],
    ["PG_PASSWORD", ".env", "root", "PostgreSQL password"],
    ["PG_DATABASE", ".env", "rag_bot_db", "PostgreSQL database name"],
    ["SECRET_KEY", ".env", "(hardcoded default)", "JWT signing key"],
    ["JWT_ALGORITHM", ".env", "HS256", "JWT algorithm"],
    ["X_API_KEY", ".env", "(hardcoded default)", "API gateway key"],
    ["TOKEN_EXPIRY_HOURS", ".env", "24", "Default JWT expiry in hours"],
    ["MEDIA_PATH", ".env", "media/", "File upload directory"],
])

doc.add_heading("11.2 RAG Configuration", level=2)
tbl(["Constant", "File", "Value", "Purpose"], [
    ["LLM_MODEL", "chatbot/config.py", "qwen3:4b", "Ollama LLM model"],
    ["EMBEDDING_MODEL", "chatbot/config.py", "nomic-embed-text", "Ollama embedding model"],
    ["OLLAMA_TIMEOUT_SECONDS", "chatbot/config.py", "300", "LLM request timeout"],
    ["VECTOR_TOP_K", "chatbot/config.py", "10", "ChromaDB retrieval count"],
    ["BM25_TOP_K", "chatbot/config.py", "10", "BM25 retrieval count"],
    ["FINAL_TOP_K", "chatbot/config.py", "5", "Reranker output count"],
    ["chunk_size", "document_loader.py", "1000", "Text splitter chunk size"],
    ["chunk_overlap", "document_loader.py", "200", "Text splitter overlap"],
])

doc.add_heading("11.3 Frontend Configuration", level=2)
tbl(["Item", "File", "Value"], [
    ["Next.js version", "package.json", "14.2.18"],
    ["React version", "package.json", "18.3.1"],
    ["react-markdown", "package.json", "10.1.0"],
    ["Backend proxy", "next.config.mjs", "/users/:path* -> http://localhost:8000"],
    ["Media proxy", "next.config.mjs", "/media/:path* -> http://localhost:8000"],
    ["API key env var", "NEXT_PUBLIC_API_KEY", "Required in .env.local"],
    ["TypeScript strict", "tsconfig.json", "true"],
    ["Path alias", "tsconfig.json", "@/* -> ./*"],
])

doc.add_heading("11.4 Database Configuration", level=2)
tbl(["Parameter", "Value", "File"], [
    ["pool_size", "5", "db/db_session.py"],
    ["max_overflow", "10", "db/db_session.py"],
    ["pool_pre_ping", "True", "db/db_session.py"],
    ["autocommit", "False", "db/db_session.py"],
    ["autoflush", "False", "db/db_session.py"],
])

# SECTION 12
doc.add_heading("12. Dependencies & Versions", level=1)

doc.add_heading("12.1 Backend Python Dependencies", level=2)
tbl(["Package", "Version", "Purpose"], [
    ["fastapi", "0.138.2", "Web framework"],
    ["uvicorn", "0.38.2", "ASGI server"],
    ["sqlalchemy", "2.1.6", "ORM / database"],
    ["psycopg2-binary", "2.9.14", "PostgreSQL driver"],
    ["langchain", "1.0.12", "LLM orchestration framework"],
    ["langchain-core", "1.3.3", "Core LangChain abstractions"],
    ["langchain-community", "1.0.12", "Community integrations"],
    ["langchain-ollama", "1.2.4", "Ollama integration"],
    ["langchain-chroma", "0.3.2", "ChromaDB integration"],
    ["langgraph", "1.1.2", "Multi-agent graph framework"],
    ["langsmith", "0.4.12", "LangChain tracing"],
    ["chromadb", "1.2.2", "Vector database"],
    ["PyPDF", "6.3.1", "PDF parsing"],
    ["sentence-transformers", "5.1.0", "Embedding models"],
    ["pyjwt", "2.14.0", "JWT token handling"],
    ["passlib", "1.7.4", "Password hashing"],
    ["argon2-cffi", "25.1.0", "Argon2 password hashing"],
    ["python-dotenv", "1.1.0", ".env file loading"],
    ["python-multipart", "0.0.31", "Form data parsing"],
    ["pydantic", "2.15.0", "Data validation"],
    ["pydantic-settings", "2.10.1", "Settings management"],
    ["pydantic-core", "2.46.1", "Pydantic core engine"],
    ["alembic", "1.18.2", "Database migrations"],
    ["bcrypt", "4.3.0", "Legacy hashing (not primary)"],
    ["pypdfium2", "4.33.0", "PDF rendering"],
    ["tiktoken", "0.11.0", "Token counting"],
    ["numpy", "2.4.3", "Numerical operations"],
    ["httpx", "0.28.1", "HTTP client"],
    ["requests", "2.32.4", "HTTP client"],
    ["pillow", "12.1.0", "Image processing (profile pictures)"],
    ["aiofiles", "24.1.0", "Async file operations"],
    ["python-dateutil", "2.9.0", "Date parsing"],
    ["regex", "2027.1.4", "Advanced regex"],
    ["tenacity", "11.1.2", "Retry logic"],
    ["packaging", "25.0", "Version parsing"],
    ["certifi", "2025.8.3", "SSL certificates"],
    ["charset-normalizer", "3.4.3", "Charset detection"],
    ["click", "8.3.2", "CLI framework"],
    ["Jinja2", "3.1.6", "Templating"],
    ["jiter", "0.11.5", "JSON parsing"],
    ["jsonpath-ng", "1.9.0", "JSON path queries"],
    ["jsonschema", "4.25.0", "JSON validation"],
    ["langsmith[openai]", "0.4.12", "LangSmith with OpenAI support"],
    ["mmh3", "5.2", "MurmurHash3"],
    ["opentelemetry-api", "1.37.0", "Tracing API"],
    ["opentelemetry-sdk", "1.37.0", "Tracing SDK"],
    ["posthog", "5.5.0", "Analytics"],
    ["pydantic-extra-types", "2.24.2", "Additional Pydantic types"],
    ["pymupdf", "1.26.3", "PDF processing"],
    ["tokenizers", "0.21.4", "Fast tokenization"],
    ["tqdm", "4.67.1", "Progress bars"],
    ["typing-extensions", "4.15.0", "Type hint backports"],
    ["urllib3", "2.5.0", "HTTP client"],
    ["win32-setctime", "1.2.2", "Windows file timestamps"],
])

doc.add_heading("12.2 Frontend Node Dependencies", level=2)
tbl(["Package", "Version", "Purpose"], [
    ["next", "^14.2.18", "React framework (App Router)"],
    ["react", "^18.3.1", "UI library"],
    ["react-dom", "^18.3.1", "React DOM rendering"],
    ["react-markdown", "^10.1.0", "Markdown rendering in chat"],
])

# SECTION 13
doc.add_heading("13. API Reference", level=1)

doc.add_heading("13.1 Authentication Endpoints", level=2)
tbl(["Method", "Endpoint", "Auth", "Request", "Response"], [
    ["POST", "/users/api/signup", "X-API-Key", "Form: display_name, email, password", "{success, status, message, data: {user_id, email, message}}"],
    ["POST", "/users/api/login", "X-API-Key", "JSON: email, password, remember_me", "{success, status, message, data: {access_token, token_type, user: {id, email, display_name, dob, profile_picture_url}}}"],
])

doc.add_heading("13.2 Profile Endpoints", level=2)
tbl(["Method", "Endpoint", "Auth", "Request", "Response"], [
    ["GET", "/users/profile-data", "JWT Bearer", "None", "{success, status, message, data: {email, display_name, dob, profile_picture_url}}"],
    ["PATCH", "/users/update-profile", "JWT Bearer", "Form: display_name?, dob?, profile_picture? (file)", "{success, status, message, data: ProfileOut}"],
])

doc.add_heading("13.3 Session Endpoints", level=2)
tbl(["Method", "Endpoint", "Auth", "Request", "Response"], [
    ["POST", "/users/api/chat/sessions", "JWT Bearer", "JSON: {domain_key?} (auto-generates title)", "{success, data: {id, title, domain_key, created_at}}"],
    ["GET", "/users/api/chat/sessions", "JWT Bearer", "None", "{success, data: [{id, title, domain_key, is_active, created_at, updated_at, message_count}]}"],
    ["PATCH", "/users/api/chat/sessions/{session_id}", "JWT Bearer", "JSON: {title: string}", "{success, data: {id, title, ...}}"],
    ["DELETE", "/users/api/chat/sessions/{session_id}", "JWT Bearer", "None", "{success, message: 'Session deleted successfully'}"],
])

doc.add_heading("13.4 Message Endpoints", level=2)
tbl(["Method", "Endpoint", "Auth", "Request", "Response"], [
    ["POST", "/users/api/chat/sessions/{session_id}/messages", "JWT Bearer", "JSON: {content: string}", "{success, data: {user_message: {id, role, content, created_at}, assistant_message: {id, role, content, sources, created_at}}}"],
    ["GET", "/users/api/chat/sessions/{session_id}/messages", "JWT Bearer", "None", "{success, data: [{id, role, content, sources, created_at}]}"],
])

doc.add_heading("13.5 Unmounted Endpoints (Not in Production)", level=2)
tbl(["Method", "Endpoint", "File", "Status"], [
    ["POST", "/chatbot/chat", "chatbot/router.py", "NOT MOUNTED in main.py"],
    ["POST", "/chat/send", "chatbot/router.py", "NOT MOUNTED in main.py"],
    ["POST", "/users/api/chat/send", "chat/routers/chat.py", "NOT MOUNTED in main.py"],
])

# SECTION 14
doc.add_heading("14. Error Handling", level=1)

doc.add_heading("14.1 Backend Error Handling", level=2)
doc.add_paragraph("Global exception handler (main.py): Catches all unhandled exceptions, returns JSONResponse {success: False, status: 'error', message: 'An unexpected error occurred', error: str(e)} with status 500.")
doc.add_paragraph("Per-route: session_controller returns HTTPException(404) if session not found or user mismatch. JWT auth returns 401 on invalid/expired token. X-API-Key returns 403 on mismatch.")
doc.add_paragraph("RAG pipeline: ChatbotService.ask() catches exceptions and returns {success: False, answer: error_message, sources: empty}. Reranker falls back to original order on JSON parse failure.")

doc.add_heading("14.2 Frontend Error Handling", level=2)
doc.add_paragraph("API errors: apiErrorMessage(err) in lib/api.ts extracts error message from response body, falls back to err.message or 'Request failed'. 403 errors trigger isForbiddenApiKeyError() check and show ApiKeySetupNotice.")
doc.add_paragraph("Chat errors: error state in dashboard/page.tsx. Failed sends result in placeholder messages being removed from optimistic update. Error displayed in UI as ChatBubble with error styling.")
doc.add_paragraph("Auth errors: jwtAuthCheck() in AppShell checks localStorage token on mount. If invalid/expired, clears session and redirects to /login.")

doc.add_heading("14.3 Error Response Format", level=2)
code("""{
  "success": false,
  "status": "error",
  "message": "Human-readable error message",
  "error": "Technical error detail (optional)"
}""")

# SECTION 15
doc.add_heading("15. Known Issues & Stubs", level=1)

doc.add_heading("15.1 Broken Components", level=2)
tbl(["Component", "File", "Issue", "Status"], [
    ["MemoryManager", "chatbot/memory/memory_manager.py", "Calls nonexistent load_recent_messages() (actual: load()); imports nonexistent PromptContextBuilder (actual: PromptContext)", "NOT USED - session_controller uses HistoryLoader directly"],
    ["ConversationSummarizer", "chatbot/memory/conversation_summary.py", "summarize() method returns None (stub)", "NOT USED"],
    ["WebAgent", "chatbot/agents/web_agent.py", "search() returns 'Web Search is not implemented yet.' (stub)", "NOT USED"],
])

doc.add_heading("15.2 Unmounted Routers", level=2)
tbl(["Router", "File", "Mounted In", "Status"], [
    ["chatbot_router", "chatbot/router.py", "main.py", "NOT MOUNTED - defines /chatbot/chat and /chat/send endpoints"],
    ["chat router", "chat/routers/chat.py", "main.py", "NOT MOUNTED - defines /users/api/chat/send endpoint"],
])

doc.add_heading("15.3 Unused Modules", level=2)
tbl(["Module", "File", "Status"], [
    ["chat_controller", "chat/controllers/chat_controller.py", "NOT USED by main flow (standalone function)"],
    ["HistorySelector", "chatbot/memory/history_selector.py", "NOT USED (functionality handled in PromptContext)"],
    ["LangGraph graph", "chatbot/graph/graph.py", "DEFINED but NOT WIRED into session_controller flow"],
    ["supervisor_node", "chatbot/nodes/supervisor_node.py", "DEFINED but NOT WIRED"],
    ["rag_node", "chatbot/nodes/rag_node.py", "DEFINED but NOT WIRED"],
    ["web_node", "chatbot/nodes/web_node.py", "DEFINED but NOT WIRED"],
])

doc.add_heading("15.4 Duplicate/Redundant Files", level=2)
tbl(["File", "Purpose", "Status"], [
    ["chat/controllers/chat_controller.py", "Standalone chat endpoint", "Duplicate of session_controller.send_chat_message() logic"],
    ["chat/routers/chat.py", "Alternative chat router", "Duplicate of messages.py router"],
    ["chat/schemas/chat_schema.py", "Standalone ChatRequest/ChatResponse schemas", "Only used by unmounted chat router"],
])

doc.add_heading("15.5 Missing Items", level=2)
doc.add_paragraph("Alembic migration files: alembic.ini and alembic/ directory configured but no explicit migration files found in codebase.")
doc.add_paragraph("LangSmith tracing: langsmith package in requirements.txt, LANGCHAIN_TRACING_V2 env var referenced in test_graph.py, but not configured in main flow.")
doc.add_paragraph("File upload validation: profile picture upload in update_user_profile.py accepts any file; no file type/size validation.")
doc.add_paragraph("Rate limiting: No rate limiting configured on any endpoint.")
doc.add_paragraph("Unit tests: No test files found in the codebase.")

# SECTION 16
doc.add_heading("16. Data Flow Diagrams", level=1)

doc.add_heading("16.1 End-to-End RAG Data Flow", level=2)
diag("End-to-End Data Flow", [
    "+-----------+     +---------+     +----------+     +--------+     +------+",
    "| PDF File  | --> | Ingest  | --> | ChromaDB |     | BM25   |     | Ollama |",
    "| (Policy)  |     | Script  |     | (Vectors)|     | Index  |     | LLM   |",
    "+-----------+     +---------+     +----+-----+     +---+----+     +---+---+",
    "                                                     |               |",
    "                                                     |               |",
    "User Query  ------------------------->  QueryRewriter -----+        |",
    "                                                     |               |",
    "                                                     v               |",
    "                                              HybridRetriever       |",
    "                                                     |               |",
    "                                                     v               |",
    "                                                Reranker ---------->|",
    "                                                     |               |",
    "                                                     v               |",
    "                                              Top 5 Documents       |",
    "                                                     |               |",
    "                                                     v               |",
    "ChatbotService.ask() <--- HistoryLoader (DB)   SYSTEM_PROMPT       |",
    "       |                                                  |          |",
    "       +------>  prompt | llm  <----- question + context + history  |",
    "       |                          (qwen3:4b)                        |",
    "       v                                                            |",
    "Answer + Sources ----> session_controller ----> DB (chat_messages)  |",
    "                            |                                       |",
    "                            v                                       |",
    "                     Frontend (ChatBubble + SourcePanel)            |",
])

doc.add_heading("16.2 Authentication Data Flow", level=2)
diag("Authentication Flow", [
    "SIGNUP:",
    "  User -> /users/api/signup (Form + X-API-Key) -> validate_password()",
    "       -> get_user_by_email() check exists -> hash_password()",
    "       -> db.add(UserRef) + db.add(UserProfile) -> commit",
    "",
    "LOGIN:",
    "  User -> /users/api/login (JSON + X-API-Key) -> get_user_by_email()",
    "       -> verify_password() (argon2) -> create_jwt_token()",
    "       -> return {access_token, user data}",
    "",
    "PROTECTED REQUEST:",
    "  User -> Authorization: Bearer {token} -> jwt_auth() -> HTTPBearer",
    "       -> decode_jwt_token() -> request.state.decoded_token",
    "       -> Controller extracts user_id -> business logic",
])

doc.add_heading("16.3 Session Lifecycle", level=2)
diag("Session Lifecycle", [
    "CREATE:",
    "  Frontend POST /users/api/chat/sessions -> create_session(user_id, domain_key?)",
    "  -> Returns session_id + title ('New Chat')",
    "",
    "SEND MESSAGE:",
    "  Frontend POST /users/api/chat/sessions/{id}/messages",
    "  -> get_session() verify ownership",
    "  -> create_message(role='user')",
    "  -> _derive_title() if title is 'New Chat' (truncate to 50 chars)",
    "  -> HistoryLoader.load() fetch last 10 messages",
    "  -> ChatbotService.ask() generate answer",
    "  -> create_message(role='assistant', sources=sources_json)",
    "  -> Return both messages to frontend",
    "",
    "LOAD HISTORY:",
    "  Frontend GET /users/api/chat/sessions/{id}/messages",
    "  -> get_messages(session_id) ordered by created_at ASC",
    "  -> Returns all messages with sources for rendering",
    "",
    "RENAME:",
    "  Frontend PATCH /users/api/chat/sessions/{id} {title: 'Leave Policy Questions'}",
    "  -> update_session(session_id, user_id, title=title)",
    "",
    "DELETE:",
    "  Frontend DELETE /users/api/chat/sessions/{id}",
    "  -> delete_session(session_id, user_id) -> sets is_deleted=True",
    "  -> Cascade deletes all chat_messages",
])

# SECTION 17
doc.add_heading("17. Conclusion", level=1)

doc.add_paragraph("This document captures the complete technical architecture of the RAG Chatbot project as it exists in the codebase. The system is a document-grounded conversational AI with a FastAPI + PostgreSQL backend, Next.js 14 frontend, and a RAG pipeline powered by LangChain, ChromaDB, BM25, and Ollama (qwen3:4b).")

doc.add_paragraph("The architecture demonstrates a well-structured modular design across 40+ Python files and 20+ TypeScript files, with clear separation of concerns across users, chat, retrieval, memory, and configuration modules. Authentication uses a dual-layer approach (API key gateway + JWT bearer) with Argon2 password hashing.")

doc.add_paragraph("Key findings from this audit:")
doc.add_paragraph("  - The core RAG pipeline (vector + BM25 hybrid search with LLM reranking) is fully functional")
doc.add_paragraph("  - The multi-agent LangGraph architecture is defined but not yet integrated into the main flow")
doc.add_paragraph("  - MemoryManager is broken and unused; HistoryLoader handles session history directly")
doc.add_paragraph("  - WebAgent is a stub returning 'not implemented'")
doc.add_paragraph("  - ConversationSummarizer is a stub returning None")
doc.add_paragraph("  - Two chat routers exist but only one is mounted in main.py")
doc.add_paragraph("  - No test suite, no Alembic migration files, no rate limiting")

doc.add_paragraph("The project is functional for its core use case: querying company policy documents via a conversational interface with source attribution. The multi-agent and web search capabilities are architecturally prepared but not yet wired into the production flow.")

# ============================================================
# SAVE
# ============================================================
save()
