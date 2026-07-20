"""
Generates the RAG Chatbot Technical Architecture Report as a .docx file.
Run: python generate_report.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime, os

doc = Document()

# ── Global style tweaks ──────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── Helper functions ──────────────────────────────────────────────────────────
def add_code_block(text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = p._element.get_or_add_pPr()
    bg = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F5F5F5",
    })
    shading.append(bg)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def add_ascii_diagram(title: str, lines: list[str]):
    doc.add_paragraph(title, style="Heading 3")
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RAG Chatbot\nTechnical Architecture Report")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("DotSquares AI — Multi-Agent RAG Assistant")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Generated: {datetime.date.today().strftime('%d %B %Y')}\n").font.size = Pt(11)
meta.add_run("Version: 0.1.0\n").font.size = Pt(11)
meta.add_run("Status: Current Codebase Snapshot").font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual placeholder)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. System Overview",
    "2. Tech Stack & Dependencies",
    "3. Directory Structure",
    "4. Backend Architecture",
    "5. Database Design",
    "6. Authentication & Authorization",
    "7. Chatbot / RAG Pipeline",
    "8. Retrieval Pipeline (Query Rewrite → Hybrid → Rerank)",
    "9. Multi-Agent Architecture (LangGraph)",
    "10. Conversation Memory",
    "11. Frontend Architecture",
    "12. API Endpoints",
    "13. Configuration & Environment",
    "14. Data Ingestion",
    "15. Error Handling & Logging",
    "16. Deployment Notes",
    "17. Summary of Gaps / Not-Present Items",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("1. System Overview", level=1)
doc.add_paragraph(
    "The RAG Chatbot is a full-stack application that answers questions about DotSquares "
    "company policies by retrieving relevant passages from a PDF document stored in a "
    "vector database and generating natural-language answers using a locally-hosted LLM."
)

doc.add_paragraph(
    "The system consists of a Python/FastAPI backend and a Next.js/React frontend. "
    "The backend exposes REST APIs for user authentication, chat session management, and "
    "the RAG question-answering pipeline. The frontend provides a chat UI with session "
    "history, suggested prompts, and source attribution."
)

add_ascii_diagram("High-Level Architecture", [
    "┌──────────────┐      REST / JSON       ┌────────────────────────────┐",
    "│              │  ◄──────────────────►   │   FastAPI Backend           │",
    "│   Next.js    │                         │   ├─ Auth (JWT + API Key)   │",
    "│   Frontend   │                         │   ├─ Session / Message CRUD │",
    "│  (port 3000) │                         │   ├─ RAG Pipeline           │",
    "│              │                         │   │   ├─ QueryRewriter       │",
    "└──────────────┘                         │   │   ├─ HybridRetriever     │",
    "                                         │   │   └─ Reranker            │",
    "                                         │   ├─ LangGraph Multi-Agent  │",
    "                                         │   └─ ChatbotService (LLM)   │",
    "                                         └────────────┬───────────────┘",
    "                                                      │",
    "                         ┌────────────────────────────┼──────────────────┐",
    "                         │                            │                  │",
    "                    ┌────▼─────┐  ┌──────────────┐  ┌─▼────────────┐",
    "                    │PostgreSQL│  │  ChromaDB     │  │   Ollama     │",
    "                    │  (Users, │  │  (Vector      │  │  qwen3:4b    │",
    "                    │ Sessions,│  │   Store)      │  │  nomic-embed │",
    "                    │ Messages)│  └──────────────┘  └──────────────┘",
    "                    └──────────┘",
])

# ═══════════════════════════════════════════════════════════════════════════════
#  2. TECH STACK & DEPENDENCIES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Tech Stack & Dependencies", level=1)

doc.add_heading("2.1 Backend", level=2)
add_table(
    ["Category", "Technology", "Version / Notes"],
    [
        ["Runtime", "Python", "3.12+ (based on environment)"],
        ["Web Framework", "FastAPI", "0.138.2"],
        ["ASGI Server", "Uvicorn", "0.49.0"],
        ["ORM", "SQLAlchemy", "2.0.51"],
        ["Migrations", "Alembic", "1.18.5"],
        ["Database", "PostgreSQL", "via psycopg2-binary 2.9.12"],
        ["LLM Framework", "LangChain", "1.3.11 (langchain-core 1.4.8)"],
        ["Agent Framework", "LangGraph", "1.2.7"],
        ["LLM Runtime", "Ollama", "ollama 0.6.2"],
        ["LLM Model", "qwen3:4b", "Local, temperature=0"],
        ["Embedding Model", "nomic-embed-text", "Local, via Ollama"],
        ["Vector Store", "ChromaDB", "1.5.9 (langchain-chroma 1.1.0)"],
        ["BM25", "rank-bm25", "0.2.2"],
        ["PDF Loader", "PyPDFLoader", "via langchain-community"],
        ["Text Splitter", "RecursiveCharacterTextSplitter", "chunk_size=1000, overlap=200"],
        ["Auth", "PyJWT", "2.13.0"],
        ["Password Hashing", "bcrypt", "5.0.0"],
        ["Validation", "Pydantic", "2.13.4"],
        ["Email Validation", "email-validator", "2.3.0"],
        ["HTTP Client", "httpx", "0.28.1"],
        ["Config", "python-dotenv", "1.2.2"],
        ["Packaging", "requirements.txt", "162 pinned dependencies"],
    ],
)

doc.add_heading("2.2 Frontend", level=2)
add_table(
    ["Category", "Technology", "Version"],
    [
        ["Framework", "Next.js", "^14.2.18"],
        ["UI Library", "React", "^18.3.1"],
        ["Language", "TypeScript", "~5.6.2"],
        ["Markdown", "react-markdown", "^10.1.0"],
        ["Styling", "CSS (globals.css)", "No Tailwind/other framework"],
        ["Fonts", "DM Sans + Outfit", "via next/font/google"],
        ["Build", "Next.js", "next build"],
        ["Dev Server", "Next.js Dev", "port 3000"],
        ["Node.js", ">=18.17.0", "checked by scripts/check-node.cjs"],
    ],
)

# ═══════════════════════════════════════════════════════════════════════════════
#  3. DIRECTORY STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Directory Structure", level=1)

add_code_block("""rag_chatbot/
├── backend/
│   ├── main.py                          # FastAPI entry point
│   ├── requirements.txt                 # 162 pinned Python dependencies
│   ├── alembic.ini                      # Alembic migration config
│   ├── alembic/                         # DB migration scripts
│   │   ├── env.py
│   │   └── versions/                    # 4 migration files
│   ├── config/
│   │   └── environment.py               # Env vars (DB URL, JWT, API key)
│   ├── db/
│   │   ├── __init__.py                  # Re-exports Base, engine, get_db
│   │   └── db_session.py               # SQLAlchemy engine, SessionLocal, get_db
│   ├── shared/
│   │   ├── auth/
│   │   │   ├── create_access_token.py   # JWT create/verify/decode
│   │   │   ├── jwt_auth_function.py     # FastAPI Bearer dependency
│   │   │   └── x_api_auth.py           # X-API-Key header dependency
│   │   ├── utils/
│   │   │   └── custom_responses.py      # Standardized JSON responses
│   │   ├── logging/                     # Logger utilities
│   │   └── debug_log.py
│   ├── src/routes/
│   │   ├── users/
│   │   │   ├── router.py               # Users APIRouter (mounted at /users)
│   │   │   ├── models/
│   │   │   │   ├── user_ref.py          # UserRef ORM (users table)
│   │   │   │   └── user_profiles.py     # UserProfile ORM (user_profile table)
│   │   │   ├── schemas/
│   │   │   │   ├── login_request.py
│   │   │   │   └── profile_out.py
│   │   │   ├── controllers/
│   │   │   │   ├── hash_password.py
│   │   │   │   ├── verify_password.py
│   │   │   │   ├── validate_password.py
│   │   │   │   └── get_user_by_email.py
│   │   │   └── routers/
│   │   │       ├── signup.py            # POST /users/signup
│   │   │       ├── login.py             # POST /users/login
│   │   │       ├── get_user_profile.py
│   │   │       └── update_user_profile.py
│   │   └── chatbot/
│   │       ├── router.py                # Chatbot APIRouter (mounted at /users)
│   │       ├── models/
│   │       │   ├── chat_session.py      # ChatSession ORM
│   │       │   └── chat_message.py      # ChatMessage ORM (JSONB sources)
│   │       ├── schemas/
│   │       │   ├── chat_request.py
│   │       │   ├── chat_response.py
│   │       │   ├── session_request.py
│   │       │   └── session_response.py
│   │       ├── crud/
│   │       │   ├── chat_session.py      # Session CRUD + title derivation
│   │       │   └── chat_message.py      # Message CRUD
│   │       ├── controllers/
│   │       │   ├── session_controller.py # send_chat_message() — main orchestration
│   │       │   └── chat_controller.py
│   │       └── routers/
│   │           ├── sessions.py          # POST/GET/PATCH/DELETE sessions
│   │           ├── messages.py          # POST/GET messages
│   │           └── chat.py
│   └── chatbot/                         # RAG engine (standalone package)
│       ├── config.py                    # LLM_MODEL, EMBEDDING_MODEL, TOP_K
│       ├── prompts.py                   # SYSTEM_PROMPT template
│       ├── chatbot_service.py           # Singleton orchestrator (ask method)
│       ├── vector_store.py             # ChromaDB singleton
│       ├── bm25_store.py              # BM25Okapi retriever
│       ├── hybrid_retriever.py         # Merges vector + BM25 with dedup
│       ├── tokenizer.py               # Regex tokenizer for BM25
│       ├── document_loader.py          # PDF load + text split
│       ├── ingest.py                   # Standalone ingestion script
│       ├── router.py                   # Standalone /chatbot/chat (NOT mounted)
│       ├── retrieval/
│       │   ├── pipeline.py             # RetrievalPipeline (rewrite→hybrid→rerank)
│       │   ├── query_rewriter.py       # Heuristic + LLM query rewriting
│       │   └── reranker.py            # LLM-based reranker
│       ├── agents/
│       │   ├── supervisor.py           # Keyword-based routing
│       │   ├── rag_agent.py            # Wraps ChatbotService.ask()
│       │   └── web_agent.py            # Stub ("not implemented")
│       ├── graph/
│       │   ├── state.py                # GraphState TypedDict
│       │   └── graph.py               # LangGraph StateGraph (NOT wired in)
│       ├── nodes/
│       │   ├── supervisor_node.py
│       │   ├── rag_node.py
│       │   └── web_node.py
│       └── memory/
│           ├── history_loader.py       # Loads last 10 messages from DB
│           ├── history_selector.py     # Takes last 6 messages
│           ├── memory_manager.py       # Orchestrator (broken — calls nonexistent method)
│           ├── conversation_summary.py # Stub returning None
│           └── prompt_context.py       # Builds "Role: Content" text
│
├── frontend/
│   ├── package.json
│   ├── app/
│   │   ├── layout.tsx                  # Root layout (DM Sans + Outfit fonts)
│   │   ├── page.tsx                    # Root page (redirects)
│   │   ├── login/page.tsx             # Login form
│   │   ├── signup/page.tsx            # Signup form
│   │   └── (app)/
│   │       ├── layout.tsx             # Authenticated layout (AppShell wrapper)
│   │       ├── dashboard/page.tsx     # Main chat interface
│   │       ├── history/page.tsx
│   │       └── profile/page.tsx
│   ├── components/
│   │   ├── AppShell.tsx               # Auth guard + collapsible sidebar
│   │   ├── ApiKeySetupNotice.tsx
│   │   ├── DotSquaresBrandLogo.tsx
│   │   ├── PasswordInput.tsx
│   │   └── chat/
│   │       ├── ChatBubble.tsx          # Message bubble (Markdown rendering)
│   │       ├── ChatComposer.tsx        # Textarea + send button
│   │       ├── AgentStepTimeline.tsx   # Multi-agent step visualization
│   │       └── SourcePanel.tsx         # Source chunk attribution panel
│   └── lib/
│       ├── api.ts                      # HTTP client, error handling
│       ├── chat-api.ts                 # Chat API functions
│       ├── auth-storage.ts             # localStorage auth management
│       └── types/
│           └── chat.ts                 # TypeScript types + formatters
""")

# ═══════════════════════════════════════════════════════════════════════════════
#  4. BACKEND ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Backend Architecture", level=1)

doc.add_heading("4.1 Entry Point — main.py", level=2)
doc.add_paragraph(
    "File: backend/main.py\n"
    "Framework: FastAPI (v0.138.2)\n"
    "CORS: Allow all origins, methods, headers (allow_origins=['*'])\n"
    "Static files: /media mounted to MEDIA_PATH (default: 'media')\n"
    "Global exception handler: catches all unhandled exceptions, returns JSON {success, status, message, error}"
)
doc.add_paragraph("Mounted routers:")
add_code_block(
    "app.include_router(users_router,  prefix='/users')\n"
    "app.include_router(chatbot_router, prefix='/users')"
)
doc.add_paragraph(
    "The chatbot_router (src/routes/chatbot/router.py) nests sessions_router and messages_router "
    "under /api/chat, resulting in final paths: /users/api/chat/sessions, /users/api/chat/sessions/{id}/messages."
)

doc.add_heading("4.2 Request Flow", level=2)
add_ascii_diagram("Chat Request Flow", [
    "Client POST /users/api/chat/sessions/{id}/messages",
    "    │",
    "    ▼",
    "messages.py router ──► jwt_auth (Bearer token verified)",
    "    │",
    "    ▼",
    "session_controller.send_chat_message()",
    "    ├── get_session() — verify session belongs to user",
    "    ├── create_message(role='user', content=...) — persist user msg",
    "    ├── _derive_title() — auto-title 'New Chat' sessions",
    "    ├── HistoryLoader(db).load(session_id) — load last 10 messages",
    "    ├── ChatbotService().ask(question, history) — RAG pipeline",
    "    │       ├── RetrievalPipeline.retrieve(question)",
    "    │       │       ├── QueryRewriter.rewrite(question)",
    "    │       │       ├── HybridRetriever.invoke(rewritten_query)",
    "    │       │       └── Reranker.rerank(query, docs)",
    "    │       ├── PromptContext.build(history) — format history",
    "    │       └── chain.invoke({history, context, question}) — LLM",
    "    └── create_message(role='assistant', content=..., sources=...)",
    "    │",
    "    ▼",
    "Response: { user_message, assistant_message }",
])

doc.add_heading("4.3 ChatbotService (Singleton)", level=2)
doc.add_paragraph(
    "File: backend/chatbot/chatbot_service.py\n"
    "Pattern: Singleton (class-level _instance + _initialized flag)\n"
    "LLM: ChatOllama(model='qwen3:4b', temperature=0, timeout=300s)\n"
    "Prompt: ChatPromptTemplate.from_template(SYSTEM_PROMPT) with {history}, {context}, {question}\n"
    "Chain: prompt | llm (LangChain LCEL)\n"
    "Method: ask(question, history=None) → dict with keys: success, answer, sources"
)
doc.add_paragraph("The ask() method:")
doc.add_paragraph("1. Calls RetrievalPipeline.retrieve(question) to get relevant documents", style="List Number")
doc.add_paragraph("2. If no docs found, returns fallback answer", style="List Number")
doc.add_paragraph("3. Formats history via PromptContext.build(history)", style="List Number")
doc.add_paragraph("4. Joins document page_content into a single context string", style="List Number")
doc.add_paragraph("5. Invokes the LLM chain with {history, context, question}", style="List Number")
doc.add_paragraph("6. Returns answer text plus source chunks with page/source metadata", style="List Number")

# ═══════════════════════════════════════════════════════════════════════════════
#  5. DATABASE DESIGN
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Database Design", level=1)

doc.add_paragraph("Database: PostgreSQL (default: rag_chatbot on localhost:5432)")
doc.add_paragraph("ORM: SQLAlchemy 2.0.51 with declarative_base()")
doc.add_paragraph("Migrations: Alembic (4 migration files in alembic/versions/)")
doc.add_paragraph(
    "Connection: pool_size=5, max_overflow=10, pool_pre_ping=True"
)

doc.add_heading("5.1 Table: users", level=2)
add_table(
    ["Column", "Type", "Constraints"],
    [
        ["id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT"],
        ["email", "TEXT", "UNIQUE, NOT NULL"],
        ["password_hash", "TEXT", "NULLABLE"],
        ["display_name", "TEXT", "NULLABLE"],
        ["dob", "DATE", "NULLABLE"],
        ["created_by_admin", "BOOLEAN", "DEFAULT FALSE, NOT NULL"],
        ["status", "TEXT", "DEFAULT 'pending', NOT NULL"],
        ["email_verified", "BOOLEAN", "DEFAULT FALSE"],
        ["otp", "TEXT", "NULLABLE"],
        ["last_login", "TIMESTAMP(tz)", "NULLABLE"],
        ["otp_created_at", "TIMESTAMP(tz)", "NULLABLE"],
        ["is_active", "BOOLEAN", "DEFAULT TRUE, NOT NULL"],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE, NOT NULL"],
        ["created_at", "TIMESTAMP(tz)", "SERVER DEFAULT now(), NOT NULL"],
        ["updated_at", "TIMESTAMP(tz)", "SERVER DEFAULT now(), ON UPDATE now()"],
        ["roles", "TEXT", "NULLABLE (comma-separated, e.g. 'user')"],
    ],
)

doc.add_heading("5.2 Table: user_profile", level=2)
add_table(
    ["Column", "Type", "Constraints"],
    [
        ["id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT"],
        ["user_id", "INTEGER", "FK → users.id ON DELETE CASCADE, UNIQUE, NOT NULL"],
        ["profile_picture_url", "TEXT", "NULLABLE"],
        ["portfolio", "JSONB", "NULLABLE"],
        ["watchlist", "JSONB", "NULLABLE"],
        ["is_active", "BOOLEAN", "DEFAULT TRUE, NOT NULL"],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE, NOT NULL"],
        ["created_at", "TIMESTAMP(tz)", "SERVER DEFAULT now(), NOT NULL"],
        ["updated_at", "TIMESTAMP(tz)", "SERVER DEFAULT now(), ON UPDATE now()"],
    ],
)

doc.add_heading("5.3 Table: chat_sessions", level=2)
add_table(
    ["Column", "Type", "Constraints"],
    [
        ["id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT"],
        ["user_id", "INTEGER", "FK → users.id, NOT NULL"],
        ["title", "TEXT", "NULLABLE"],
        ["is_active", "BOOLEAN", "DEFAULT TRUE, NOT NULL"],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE, NOT NULL"],
        ["domain_key", "TEXT", "NULLABLE"],
        ["created_at", "TIMESTAMP(tz)", "SERVER DEFAULT now(), NOT NULL"],
        ["updated_at", "TIMESTAMP(tz)", "SERVER DEFAULT now(), ON UPDATE now()"],
    ],
)

doc.add_heading("5.4 Table: chat_messages", level=2)
add_table(
    ["Column", "Type", "Constraints"],
    [
        ["id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT"],
        ["session_id", "INTEGER", "FK → chat_sessions.id, NOT NULL"],
        ["role", "TEXT", "NOT NULL ('user' or 'assistant')"],
        ["content", "TEXT", "NOT NULL"],
        ["sources", "JSONB", "NULLABLE (stores chunks, agent_steps, etc.)"],
        ["created_at", "TIMESTAMP(tz)", "SERVER DEFAULT now(), NOT NULL"],
    ],
)

doc.add_heading("5.5 Relationships", level=2)
doc.add_paragraph("UserRef (1) ──► (1) UserProfile  [back_populates, uselist=False, cascade delete]")
doc.add_paragraph("UserRef (1) ──► (N) ChatSession  [via user_id FK]")
doc.add_paragraph("ChatSession (1) ──► (N) ChatMessage  [back_populates, cascade delete-orphan]")

# ═══════════════════════════════════════════════════════════════════════════════
#  6. AUTHENTICATION & AUTHORIZATION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Authentication & Authorization", level=1)

doc.add_heading("6.1 Two-Layer Auth", level=2)
doc.add_paragraph(
    "Layer 1 — X-API-Key (x_api_auth.py): Every request must include the header "
    "X-API-Key matching the backend's X_API_KEY env var. Returns 403 on mismatch. "
    "Used by: signup, login endpoints."
)
doc.add_paragraph(
    "Layer 2 — JWT Bearer Token (jwt_auth_function.py): Protected endpoints use "
    "HTTPBearer scheme. Token is verified via PyJWT (HS256). Returns 401 on invalid/expired token. "
    "Payload contains: user_id, email, roles."
)

doc.add_heading("6.2 Token Lifecycle", level=2)
doc.add_paragraph("Creation: create_jwt_token(data, expiry_minutes, remember_me) → PyJWT encode")
doc.add_paragraph("Default expiry: TOKEN_EXPIRY_HOURS (24h). If remember_me=True: 30 days.")
doc.add_paragraph("Verification: verify_jwt_token(token) → decoded payload or None")
doc.add_paragraph("Usage: jwt_auth() FastAPI dependency extracts user data from Bearer token")

doc.add_heading("6.3 Password Handling", level=2)
doc.add_paragraph("Hashing: bcrypt (passlib + bcrypt library)")
doc.add_paragraph("Validation: validate_password() checks password requirements (controllers/validate_password.py)")
doc.add_paragraph("Email validation: email_validator library with check_deliverability=True")

doc.add_heading("6.4 Frontend Auth Storage", level=2)
doc.add_paragraph("Token: localStorage['access_token']")
doc.add_paragraph("User snapshot: localStorage['rag_chatbot_user'] → {user_id, email, display_name, roles, profile_image_url}")
doc.add_paragraph("Current session: localStorage['rag_chatbot_session_id']")
doc.add_paragraph("AppShell.tsx: Redirects to /login if no token found on mount")

# ═══════════════════════════════════════════════════════════════════════════════
#  7. CHATBOT / RAG PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Chatbot / RAG Pipeline", level=1)

doc.add_heading("7.1 Prompt Template", level=2)
doc.add_paragraph(
    "File: backend/chatbot/prompts.py\n"
    "Template variables: {history}, {context}, {question}\n\n"
    "SYSTEM_PROMPT rules:\n"
    "- Answer ONLY from retrieved context\n"
    "- Use conversation history for follow-up questions\n"
    "- Never mention PDFs, vector databases, embeddings, LangChain\n"
    "- Format every response with Markdown (headings, bullets, bold)\n"
    "- Professional, helpful, concise tone"
)

doc.add_heading("7.2 Pipeline Architecture", level=2)
add_ascii_diagram("Retrieval Pipeline", [
    "User Question",
    "    │",
    "    ▼",
    "┌──────────────────────────────┐",
    "│   QueryRewriter              │",
    "│   ├─ needs_rewrite() check   │",
    "│   │  (short query? abbrev?)  │",
    "│   └─ LLM rewrite if needed   │",
    "└──────────────┬───────────────┘",
    "               │ Rewritten Query",
    "               ▼",
    "┌──────────────────────────────┐",
    "│   HybridRetriever            │",
    "│   ├─ Vector Store (top 10)   │",
    "│   │  ChromaDB + nomic-embed  │",
    "│   ├─ BM25 Retriever (top 10) │",
    "│   │  BM25Okapi + regex tok   │",
    "│   └─ Merge + SHA-256 dedup   │",
    "└──────────────┬───────────────┘",
    "               │ Merged Documents",
    "               ▼",
    "┌──────────────────────────────┐",
    "│   Reranker (LLM-based)      │",
    "│   ├─ Send docs to LLM       │",
    "│   ├─ Get JSON index order    │",
    "│   └─ Return top FINAL_TOP_K  │",
    "│      (default 5)             │",
    "└──────────────┬───────────────┘",
    "               │ Ranked Documents",
    "               ▼",
    "┌──────────────────────────────┐",
    "│   ChatbotService.ask()       │",
    "│   ├─ Build context (join)    │",
    "│   ├─ Build history context   │",
    "│   └─ LLM chain.invoke()      │",
    "│      (qwen3:4b, temp=0)     │",
    "└──────────────────────────────┘",
])

doc.add_heading("7.3 Configuration Values", level=2)
add_table(
    ["Parameter", "Value", "Source"],
    [
        ["LLM_MODEL", "qwen3:4b", "chatbot/config.py"],
        ["EMBEDDING_MODEL", "nomic-embed-text", "chatbot/config.py"],
        ["OLLAMA_TIMEOUT_SECONDS", "300 (env override)", "chatbot/config.py"],
        ["VECTOR_TOP_K", "10", "chatbot/config.py"],
        ["BM25_TOP_K", "10", "chatbot/config.py"],
        ["FINAL_TOP_K", "5", "chatbot/config.py"],
        ["chunk_size", "1000", "document_loader.py"],
        ["chunk_overlap", "200", "document_loader.py"],
    ],
)

# ═══════════════════════════════════════════════════════════════════════════════
#  8. RETRIEVAL PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Retrieval Pipeline (Query Rewrite → Hybrid → Rerank)", level=1)

doc.add_heading("8.1 QueryRewriter", level=2)
doc.add_paragraph(
    "File: backend/chatbot/retrieval/query_rewriter.py\n"
    "Heuristic check (needs_rewrite): Returns True if query has ≤2 words, matches known abbreviations "
    "(wfh, pto, hr, pf, esi, lop, ot, salary, leave, holiday, notice period), or lacks question words.\n"
    "LLM rewrite: Sends query to qwen3:4b with instruction to expand abbreviations and optimize for "
    "company policy retrieval. Returns rewritten query string.\n"
    "If query is already clear (has question words like what/how/why/etc.), returns unchanged."
)

doc.add_heading("8.2 HybridRetriever", level=2)
doc.add_paragraph(
    "File: backend/chatbot/hybrid_retriever.py\n"
    "Vector retrieval: ChromaDB with OllamaEmbeddings(nomic-embed-text), returns top 10.\n"
    "BM25 retrieval: BM25Okapi with regex tokenizer (lowercase, strip non-alphanumeric), returns top 10.\n"
    "Merging: Concatenates vector_docs + bm25_docs, deduplicates by SHA-256 hash of page_content."
)

doc.add_heading("8.3 Reranker", level=2)
doc.add_paragraph(
    "File: backend/chatbot/retrieval/reranker.py\n"
    "Method: Sends all documents to qwen3:4b with instruction to return a JSON array of document "
    "indexes ordered by relevance.\n"
    "Fallback: If JSON parsing fails or LLM returns invalid order, falls back to original order.\n"
    "Output: Returns top FINAL_TOP_K (5) documents."
)

doc.add_heading("8.4 Vector Store", level=2)
doc.add_paragraph(
    "File: backend/chatbot/vector_store.py\n"
    "Storage: ChromaDB at backend/chatbot/chroma_db/\n"
    "Embeddings: OllamaEmbeddings(model='nomic-embed-text')\n"
    "Pattern: Module-level singleton (_vector_db global), lazy initialization\n"
    "Retriever: as_retriever(search_kwargs={'k': k})"
)

doc.add_heading("8.5 BM25 Store", level=2)
doc.add_paragraph(
    "File: backend/chatbot/bm25_store.py\n"
    "Tokenizer: Regex-based (lowercase, remove non-alphanumeric, split on whitespace)\n"
    "Pattern: Module-level singleton (_bm25 global), loads chunks from document_loader.get_chunks()\n"
    "Note: BM25 is rebuilt from scratch on every restart (not persisted)"
)

# ═══════════════════════════════════════════════════════════════════════════════
#  9. MULTI-AGENT ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Multi-Agent Architecture (LangGraph)", level=1)

doc.add_paragraph(
    "The codebase contains a LangGraph-based multi-agent architecture. "
    "NOTE: This graph is defined but is NOT currently wired into the main request flow. "
    "session_controller.py calls ChatbotService.ask() directly, not through the graph."
)

doc.add_heading("9.1 Graph State", level=2)
doc.add_paragraph(
    "File: backend/chatbot/graph/state.py\n"
    "TypedDict with fields: question, history, route, answer, sources, success, error"
)

doc.add_heading("9.2 Graph Definition", level=2)
doc.add_paragraph(
    "File: backend/chatbot/graph/graph.py\n"
    "Nodes: supervisor, rag, web\n"
    "Entry: supervisor\n"
    "Conditional edges: supervisor routes to 'rag' or 'web' based on state['route']\n"
    "Terminal: Both rag and web connect to END\n"
    "Compiled: graph = builder.compile()"
)

doc.add_heading("9.3 Agents", level=2)

doc.add_paragraph("SupervisorAgent (backend/chatbot/agents/supervisor.py):", style="List Bullet")
doc.add_paragraph(
    "Keyword-based routing. Checks if question contains: today, latest, news, current, "
    "recent, google, internet, search, 2026. Returns 'web' if found, else 'rag'."
)

doc.add_paragraph("RagAgent (backend/chatbot/agents/rag_agent.py):", style="List Bullet")
doc.add_paragraph("Wraps ChatbotService.ask(question, history). Returns the result dict.")

doc.add_paragraph("WebAgent (backend/chatbot/agents/web_agent.py):", style="List Bullet")
doc.add_paragraph("Stub. Returns 'Web Search is not implemented yet.'")

doc.add_heading("9.4 Graph Nodes", level=2)
doc.add_paragraph(
    "supervisor_node: Creates SupervisorAgent, calls route(state['question']), sets state['route']\n"
    "rag_node: Creates RagAgent, calls run(state['question']), sets state['answer']\n"
    "web_node: Creates WebAgent, calls search(state['question']), sets state['answer']\n\n"
    "All nodes are module-level singletons."
)

add_ascii_diagram("LangGraph Flow (Not Currently Active)", [
    "         ┌────────────┐",
    "         │ supervisor │",
    "         └─────┬──────┘",
    "               │",
    "        ┌──────┴──────┐",
    "        │  route check │",
    "        └──────┬──────┘",
    "          ┌────┴────┐",
    "          ▼         ▼",
    "    ┌──────────┐ ┌─────────┐",
    "    │   rag    │ │   web   │",
    "    └────┬─────┘ └────┬────┘",
    "         │             │",
    "         ▼             ▼",
    "        END           END",
])

# ═══════════════════════════════════════════════════════════════════════════════
#  10. CONVERSATION MEMORY
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Conversation Memory", level=1)

doc.add_paragraph(
    "The codebase contains a memory subsystem in chatbot/memory/. "
    "The active path uses HistoryLoader + PromptContext directly from session_controller. "
    "MemoryManager exists but is NOT used (calls nonexistent load_recent_messages method)."
)

doc.add_heading("10.1 HistoryLoader (Active)", level=2)
doc.add_paragraph(
    "File: backend/chatbot/memory/history_loader.py\n"
    "Loads last N messages (default 10) for a session from PostgreSQL via SQLAlchemy.\n"
    "Orders by created_at DESC, then reverses to chronological order.\n"
    "Returns list of {'role': str, 'content': str} dicts."
)

doc.add_heading("10.2 PromptContext (Active)", level=2)
doc.add_paragraph(
    "File: backend/chatbot/memory/prompt_context.py\n"
    "Static method build(history): Formats history list into 'Role: Content' lines.\n"
    "Returns 'No previous conversation.' if history is empty/None."
)

doc.add_heading("10.3 HistorySelector (Unused)", level=2)
doc.add_paragraph(
    "File: backend/chatbot/memory/history_selector.py\n"
    "Simple: returns messages[-max_messages:] (default 6)."
)

doc.add_heading("10.4 ConversationSummarizer (Stub)", level=2)
doc.add_paragraph(
    "File: backend/chatbot/memory/conversation_summary.py\n"
    "Returns None. Not implemented."
)

doc.add_heading("10.5 MemoryManager (Broken)", level=2)
doc.add_paragraph(
    "File: backend/chatbot/memory/memory_manager.py\n"
    "Imports PromptContextBuilder (does not exist — actual class is PromptContext).\n"
    "Calls self.loader.load_recent_messages(session_id) — method does not exist (actual method is load()).\n"
    "Not used anywhere in the codebase."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  11. FRONTEND ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Frontend Architecture", level=1)

doc.add_heading("11.1 Framework & Routing", level=2)
doc.add_paragraph(
    "Framework: Next.js 14 (App Router)\n"
    "Root layout: app/layout.tsx — sets fonts (DM Sans, Outfit), metadata\n"
    "Authenticated layout: app/(app)/layout.tsx — wraps children in AppShell\n"
    "Routes:\n"
    "  / → Root page\n"
    "  /login → Login form\n"
    "  /signup → Signup form\n"
    "  /dashboard → Main chat interface (query param: ?session={id})\n"
    "  /history → Chat history\n"
    "  /profile → User profile"
)

doc.add_heading("11.2 AppShell Component", level=2)
doc.add_paragraph(
    "File: frontend/components/AppShell.tsx\n"
    "Role: Auth guard + collapsible sidebar layout\n"
    "Features:\n"
    "- Redirects to /login if no access token in localStorage\n"
    "- Fetches user snapshot and chat sessions on mount\n"
    "- Displays session list grouped by date (Today, Yesterday, Previous 7 Days, etc.)\n"
    "- Context menu: Rename, Delete sessions\n"
    "- New Chat button: creates session via API\n"
    "- User dropdown: Profile, Log out\n"
    "- Sidebar toggle (persisted in localStorage)\n"
    "- Backfill session titles on mount"
)

doc.add_heading("11.3 Dashboard (Chat Interface)", level=2)
doc.add_paragraph(
    "File: frontend/app/(app)/dashboard/page.tsx\n"
    "Features:\n"
    "- Auto-creates chat session on first load (or restores from URL param / localStorage)\n"
    "- 6 random suggested prompts from a pool of 19 DotSquares-specific questions\n"
    "- Optimistic UI: placeholder assistant bubble while waiting for response\n"
    "- Auto-scroll to bottom on new messages\n"
    "- Auto-title: first user message becomes session title (truncated to 50 chars)\n"
    "- Error handling with user-facing messages\n"
    "- Suspense-wrapped with loading states"
)

doc.add_heading("11.4 Chat Components", level=2)
add_table(
    ["Component", "File", "Purpose"],
    [
        ["ChatBubble", "components/chat/ChatBubble.tsx",
         "Renders user/assistant messages. Assistant messages rendered with react-markdown. "
         "Includes copy button for code blocks, AgentStepTimeline, SourcePanel."],
        ["ChatComposer", "components/chat/ChatComposer.tsx",
         "Auto-resizing textarea + send button. Enter to send, Shift+Enter for newline. "
         "Loading dots animation while busy."],
        ["AgentStepTimeline", "components/chat/AgentStepTimeline.tsx",
         "Displays multi-agent pipeline steps with labels, status, detail, duration. "
         "Shows 'Working…' pulse when liveSteps=true."],
        ["SourcePanel", "components/chat/SourcePanel.tsx",
         "Collapsible panel showing source chunks with rank, score, preview, metadata. "
         "Displays rewritten query and history summary if available."],
    ],
)

doc.add_heading("11.5 API Client", level=2)
doc.add_paragraph(
    "File: frontend/lib/api.ts\n"
    "Features:\n"
    "- apiJson(path, init): Core fetch wrapper with X-API-Key and Bearer token headers\n"
    "- apiErrorMessage(ex): Extracts error text from various backend response formats\n"
    "- isForbiddenApiKeyError(ex): Detects 403 API key errors\n"
    "- unwrapApiData<T>(raw): Extracts data from {success, data} response envelopes\n"
    "- Credentials: 'same-origin', cache: 'no-store'"
)

doc.add_heading("11.6 Chat API Functions", level=2)
doc.add_paragraph(
    "File: frontend/lib/chat-api.ts\n"
    "Functions:\n"
    "- createChatSession(opts?) → POST /users/api/chat/sessions\n"
    "- listChatSessions() → GET /users/api/chat/sessions\n"
    "- getChatMessages(sessionId) → GET /users/api/chat/sessions/{id}/messages\n"
    "- sendChatMessage(sessionId, content) → POST /users/api/chat/sessions/{id}/messages\n"
    "- updateSessionTitle(sessionId, title) → PATCH /users/api/chat/sessions/{id}\n"
    "- backfillSessionTitles() → POST /users/api/chat/sessions/backfill-titles\n"
    "- deleteSession(sessionId) → DELETE /users/api/chat/sessions/{id}"
)

doc.add_heading("11.7 TypeScript Types", level=2)
doc.add_paragraph(
    "File: frontend/lib/types/chat.ts\n"
    "Types: SourceChunk, AgentStep, MessageSources, ChatSessionRow, ChatMessageRow, ChatMessage, DateGroup\n"
    "Functions: mapMessageRows(), formatMessageTime(), groupSessionsByDate()\n"
    "Constants: AGENT_LABELS (history_analyst, retriever, grader, query_optimizer, synthesizer)"
)

# ═══════════════════════════════════════════════════════════════════════════════
#  12. API ENDPOINTS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("12. API Endpoints", level=1)

doc.add_heading("12.1 Root Endpoints", level=2)
add_table(
    ["Method", "Path", "Auth", "Description"],
    [
        ["GET", "/", "None", "Health check — returns {message, status}"],
        ["GET", "/health", "None", "Returns {status: 'healthy'}"],
    ],
)

doc.add_heading("12.2 User Endpoints", level=2)
add_table(
    ["Method", "Path", "Auth", "Description"],
    [
        ["POST", "/users/signup", "X-API-Key", "Create account. Form data: display_name, email, password"],
        ["POST", "/users/login", "X-API-Key", "Login. JSON: email, password, remember_me. Returns JWT"],
        ["GET", "/users/profile", "JWT", "Get user profile"],
        ["PATCH", "/users/profile", "JWT", "Update user profile"],
    ],
)

doc.add_heading("12.3 Chatbot Endpoints", level=2)
add_table(
    ["Method", "Path", "Auth", "Description"],
    [
        ["POST", "/users/api/chat/sessions", "JWT", "Create new chat session. JSON: title?, domain_key?"],
        ["GET", "/users/api/chat/sessions", "JWT", "List all sessions (with message counts)"],
        ["PATCH", "/users/api/chat/sessions/{id}", "JWT", "Update session title"],
        ["DELETE", "/users/api/chat/sessions/{id}", "JWT", "Soft-delete session (sets is_deleted=True)"],
        ["POST", "/users/api/chat/sessions/backfill-titles", "JWT", "Auto-title 'New Chat' sessions"],
        ["GET", "/users/api/chat/sessions/{id}/messages", "JWT", "Get all messages for a session"],
        ["POST", "/users/api/chat/sessions/{id}/messages", "JWT", "Send message + get AI response. JSON: {content}"],
    ],
)

doc.add_paragraph(
    "Note: backend/chatbot/router.py defines a standalone POST /chatbot/chat endpoint but is NOT "
    "mounted in main.py."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  13. CONFIGURATION & ENVIRONMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Configuration & Environment", level=1)

doc.add_heading("13.1 Backend Environment Variables", level=2)
add_table(
    ["Variable", "Default", "Source"],
    [
        ["BASE_URL", "http://127.0.0.1:8000", "config/environment.py"],
        ["PG_HOST", "localhost", "config/environment.py"],
        ["PG_PORT", "5432", "config/environment.py"],
        ["PG_USER", "postgres", "config/environment.py"],
        ["PG_PASSWORD", "admin", "config/environment.py"],
        ["PG_DATABASE", "rag_chatbot", "config/environment.py"],
        ["SECRET_KEY", "change-me-in-production", "config/environment.py"],
        ["JWT_ALGORITHM", "HS256", "config/environment.py"],
        ["X_API_KEY", "dev-x-api-key-change-me", "config/environment.py"],
        ["TOKEN_EXPIRY_HOURS", "24", "config/environment.py"],
        ["MEDIA_PATH", "media", "config/environment.py"],
        ["APP_NAME", "RAG Chatbot", "config/environment.py"],
        ["OLLAMA_TIMEOUT_SECONDS", "300", "chatbot/config.py"],
    ],
)

doc.add_heading("13.2 Frontend Environment Variables", level=2)
add_table(
    ["Variable", "Default", "Source"],
    [
        ["NEXT_PUBLIC_API_URL", "(empty — same origin)", "frontend/lib/api.ts"],
        ["NEXT_PUBLIC_X_API_KEY", "(empty)", "frontend/lib/api.ts"],
    ],
)

doc.add_heading("13.3 Chatbot Configuration", level=2)
doc.add_paragraph(
    "File: backend/chatbot/config.py\n"
    "DATA_PATH: backend/chatbot/data/Dotsquares_Company_Policy.pdf\n"
    "CHROMA_DB_PATH: backend/chatbot/chroma_db/\n"
    "LLM_MODEL: qwen3:4b\n"
    "EMBEDDING_MODEL: nomic-embed-text\n"
    "VECTOR_TOP_K: 10 | BM25_TOP_K: 10 | FINAL_TOP_K: 5"
)

# ═══════════════════════════════════════════════════════════════════════════════
#  14. DATA INGESTION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("14. Data Ingestion", level=1)

doc.add_heading("14.1 Standalone Ingestion Script", level=2)
doc.add_paragraph(
    "File: backend/chatbot/ingest.py\n"
    "Run via: python -m chatbot.ingest\n"
    "Steps:\n"
    "1. Load PDF using PyPDFLoader(DATA_PATH)\n"
    "2. Split into chunks: RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)\n"
    "3. Create ChromaDB: Chroma.from_documents(chunks, OllamaEmbeddings, persist_directory)\n"
    "Output: backend/chatbot/chroma_db/ directory"
)

doc.add_heading("14.2 Runtime Document Loading", level=2)
doc.add_paragraph(
    "File: backend/chatbot/document_loader.py\n"
    "Used by BM25Retriever to load chunks on startup.\n"
    "load_documents(): PyPDFLoader(DATA_PATH).load()\n"
    "split_documents(): RecursiveCharacterTextSplitter(1000, 200).split_documents()\n"
    "get_chunks(): Convenience function combining both."
)
doc.add_paragraph(
    "NOTE: BM25 store loads documents from the PDF file directly (not from ChromaDB), "
    "so both vector and BM25 stores must be in sync with the same PDF."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  15. ERROR HANDLING & LOGGING
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("15. Error Handling & Logging", level=1)

doc.add_heading("15.1 Global Exception Handler", level=2)
doc.add_paragraph(
    "File: backend/main.py\n"
    "Catches all unhandled exceptions. HTTPException → returns detail as error. "
    "Other exceptions → logs full traceback, returns 500 with generic message."
)

doc.add_heading("15.2 Custom Response Utilities", level=2)
doc.add_paragraph(
    "File: backend/shared/utils/custom_responses.py\n"
    "custom_error_response(message, error, status_code) → JSONResponse with {success: False}\n"
    "custom_response(message, data, status_code) → dict with {success: True}"
)

doc.add_heading("15.3 Chatbot Error Handling", level=2)
doc.add_paragraph(
    "chatbot_service.ask(): Wrapped in try/except. Returns {success: False, answer: 'Something went wrong'} on exception.\n"
    "session_controller.send_chat_message(): On chatbot.ask() failure, creates a fallback assistant message "
    "'Sorry, something went wrong while generating the response.' and returns it."
)

doc.add_heading("15.4 Logging", level=2)
doc.add_paragraph(
    "Backend: Python logging module with DEBUG level, format: 'asctime - level - name - message'\n"
    "Chatbot service: Uses logging.getLogger(__name__) with detailed debug/error logging\n"
    "Routers: Structured logging with session_id, user_id, elapsed_ms for performance tracking"
)

doc.add_heading("15.5 Frontend Error Handling", level=2)
doc.add_paragraph(
    "api.ts: apiErrorMessage() extracts error text from various backend response formats\n"
    "isForbiddenApiKeyError(): Detects 403 API key errors for specific user messaging\n"
    "Dashboard: Catches errors from sendChatMessage, removes optimistic messages, shows error to user"
)

# ═══════════════════════════════════════════════════════════════════════════════
#  16. DEPLOYMENT NOTES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("16. Deployment Notes", level=1)

doc.add_heading("16.1 Backend", level=2)
doc.add_paragraph(
    "Server: Uvicorn ASGI server (uvicorn 0.49.0)\n"
    "Database: PostgreSQL — requires psycopg2-binary, Alembic migrations\n"
    "LLM: Ollama must be running locally with qwen3:4b and nomic-embed-text models pulled\n"
    "Vector DB: ChromaDB stored at backend/chatbot/chroma_db/ (must run ingest.py first)\n"
    "Static files: /media directory auto-created by main.py"
)

doc.add_heading("16.2 Frontend", level=2)
doc.add_paragraph(
    "Server: Next.js dev server (port 3000) or next start for production\n"
    "Node.js: >=18.17.0 (checked by scripts/check-node.cjs)\n"
    "Build: next build\n"
    "Environment: NEXT_PUBLIC_API_URL, NEXT_PUBLIC_X_API_KEY must be set"
)

doc.add_heading("16.3 Prerequisites", level=2)
doc.add_paragraph(
    "1. PostgreSQL running with 'rag_chatbot' database created\n"
    "2. Alembic migrations applied: alembic upgrade head\n"
    "3. Ollama running with models: ollama pull qwen3:4b && ollama pull nomic-embed-text\n"
    "4. PDF ingested: python -m chatbot.ingest\n"
    "5. Backend running: uvicorn main:app --reload\n"
    "6. Frontend running: npm run dev"
)

# ═══════════════════════════════════════════════════════════════════════════════
#  17. SUMMARY OF GAPS / NOT-PRESENT ITEMS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("17. Summary of Gaps / Not-Present Items", level=1)

doc.add_paragraph(
    "The following items exist in the codebase but are incomplete, broken, or not integrated:"
)

add_table(
    ["Item", "Status", "Details"],
    [
        ["LangGraph multi-agent graph",
         "Defined but NOT wired in",
         "chatbot/graph/graph.py defines a compiled StateGraph but session_controller.py "
         "calls ChatbotService.ask() directly. The graph is never imported or invoked."],
        ["WebAgent",
         "Stub",
         "Returns 'Web Search is not implemented yet.' No web search capability."],
        ["SupervisorAgent routing",
         "Not used in main flow",
         "Keyword-based routing exists but is only used by the unused graph."],
        ["MemoryManager",
         "Broken",
         "Calls nonexistent load_recent_messages() method and imports nonexistent "
         "PromptContextBuilder class. Never called by any other code."],
        ["ConversationSummarizer",
         "Stub",
         "Returns None. No conversation summarization implemented."],
        ["chatbot/router.py endpoint",
         "NOT mounted",
         "Defines POST /chatbot/chat but main.py does not include this router."],
        ["Lifespan handler",
         "Commented out",
         "main.py has a commented-out lifespan that would clear pipeline cache."],
        ["Streaming responses",
         "Not Present",
         "No SSE or WebSocket streaming. Responses are returned as complete JSON."],
        ["Rate limiting",
         "Not Present",
         "No rate limiting middleware on any endpoint."],
        ["Tests",
         "Not Present",
         "No test files found in the codebase."],
        ["Docker/containerization",
         "Not Present",
         "No Dockerfile or docker-compose.yml found."],
        ["CI/CD",
         "Not Present",
         "No CI/CD configuration files found."],
        ["Frontend state management",
         "Not Present",
         "No Redux, Zustand, or other state library. Uses React useState/localStorage."],
        ["Frontend styling framework",
         "Not Present",
         "Uses plain CSS (globals.css). No Tailwind, CSS Modules, or styled-components."],
    ],
)

# ── Save ──────────────────────────────────────────────────────────────────────
OUTPUT = "RAG_Chatbot_Technical_Architecture_Report.docx"
doc.save(OUTPUT)
print(f"Report saved: {OUTPUT}")
print(f"Full path: {os.path.abspath(OUTPUT)}")
